"""
Lecture processing pipeline: download PDF, process slides, generate outputs, upload to S3
"""

import os
import io
import json
import logging
from typing import Dict, Any, List, Optional
from openai import OpenAI
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import tempfile
import genanki

from storage import download_from_s3, upload_to_s3, generate_presigned_url

logger = logging.getLogger(__name__)

# Initialize OpenAI client
openai_client = None


def get_openai_client() -> OpenAI:
    """Lazy-initialize OpenAI client"""
    global openai_client
    if openai_client is None:
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OPENAI_API_KEY not set")
        openai_client = OpenAI(api_key=api_key)
    return openai_client


def extract_slides_from_pdf(pdf_bytes: bytes) -> List[bytes]:
    """
    Extract individual slides/pages from PDF as PNG images
    
    Args:
        pdf_bytes: PDF file bytes
    
    Returns:
        List of PNG image bytes for each slide
    """
    slides = []
    
    try:
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            
            # Convert to image (300 DPI for quality)
            mat = fitz.Matrix(300/72, 300/72)
            pix = page.get_pixmap(matrix=mat)
            
            # Convert to PNG bytes
            img_data = pix.tobytes("png")
            slides.append(img_data)
        
        pdf_document.close()
        logger.info(f"Extracted {len(slides)} slides from PDF")
        return slides
    
    except Exception as e:
        logger.error(f"Error extracting slides: {e}")
        raise


def get_prompt(language: str = "Spanish") -> str:
    """Get the prompt template for the specified language"""
    language_instruction = f"\n- Esta explicación debe ser escrita en {language}.\n"
    
    return f"""\
Hazme una explicación **completa, clara y dinámica** sobre este texto.{language_instruction}
Debe permitir al lector **entender todo el contenido técnico de manera fácil y ordenada**,
sin extenderse demasiado ni omitir ningún detalle importante.
Incluye ejemplos o analogías cuando ayuden a comprender mejor.
Al final, agrega un **resumen corto** con lo más importante de toda la explicación.

OBJETIVO GENERAL
- Que sea **profunda pero comprensible**, con rigor técnico y tono didáctico.
- Que ayude a **aprender de forma rápida**.
- Que combine explicación fluida con **puntos clave**.
- **IMPORTANTE:** Siempre genera texto explicativo completo, incluso si la diapositiva es un gráfico, diagrama o imagen sin texto. Analiza visualmente y describe lo que ves, explicando su significado y relevancia.
- **IDIOMA:** Todas las explicaciones deben estar en **{language}**, pero mantén las **palabras técnicas más importantes** (términos clave, conceptos específicos) en su **idioma original** (inglés, alemán, etc.) para facilitar el aprendizaje de vocabulario técnico.

INSTRUCCIONES
1) Explica el tema principal y por qué es relevante.
2) **EXPLICACIÓN DIDÁCTICA:** Divide la explicación completa en **puntos clave detallados y profundos** (no un párrafo largo). Cada punto debe ser **súper completo, técnico y profesional**, cubriendo TODOS los detalles visibles en la diapositiva sin omitir absolutamente nada. Explica conceptos complejos de manera que un principiante pueda entenderlos desde cero, pero con rigor técnico suficiente para convertir al lector en un experto absoluto que domine los conceptos y pueda usar términos técnicos correctamente. Incluye definiciones, ejemplos prácticos, analogías cuando ayuden, y conexiones lógicas. Mantén términos técnicos importantes en inglés o alemán si aplica, explicándolos en {language} cuando sea necesario. Usa más términos en inglés para conceptos clave y nombres específicos del PowerPoint, explicándolos en {language} cuando sea necesario. Proporciona el contenido directo sin prefijos como "Punto 1:", "Punto 2:", etc.
3) Resume conceptos principales adicionales en puntos clave (usando términos originales donde sea clave).
4) Conecta con temas relacionados, pero haciéndolo específico y en relación con las demás diapositivas, no tan general. Aporta información realmente útil y que ayude a comprender mejor el tema, no datos innecesarios.
5) Cierra con un **resumen corto** (2–3 frases con el takeaway).
6) **SIEMPRE** proporciona contenido completo, no dejes campos vacíos o con "N/A".
7) Genera 3 ankis de conceptos importantes de aprender e interiorizar en esta diapositiva.

FORMATO DE SALIDA
Devuelve **únicamente** un **objeto JSON válido** (sin texto adicional, sin comentarios).
NO copies literalmente el ejemplo; rellénalo con el contenido del slide.

```json
{{
  "titulo": "Tema o concepto central de la diapositiva",
  "explicacion_didactica": ["Punto 1: Detalle completo...", "Punto 2: Detalle completo...", "Punto 3: ..."],
  "puntos_clave": ["Idea 1", "Idea 2", "Idea 3"],
  "conexiones": "Relaciones con otros temas importantes",
  "resumen_corto": "Síntesis breve (2–3 frases)",
  "anki_cards": [
    {{"pregunta": "...", "respuesta": "..."}},
    {{"pregunta": "...", "respuesta": "..."}},
    {{"pregunta": "...", "respuesta": "..."}}
  ]
}}
"""


def explain_slide(slide_bytes: bytes, slide_number: int, language: str = "Spanish") -> Dict[str, Any]:
    """
    Generate explanation for a single slide using OpenAI Vision API
    
    Args:
        slide_bytes: PNG image bytes of the slide
        slide_number: Slide number (1-indexed)
        language: Language for explanation
    
    Returns:
        Dictionary with slide explanation
    """
    import base64
    import re
    
    client = get_openai_client()
    
    try:
        # Encode image to base64
        image_base64 = base64.b64encode(slide_bytes).decode('utf-8')
        image_url = f"data:image/png;base64,{image_base64}"
        
        prompt = get_prompt(language)
        
        # Call Vision API
        response = client.chat.completions.create(
            model="gpt-4o",
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": image_url, "detail": "high"}}
                    ]
                }
            ],
            max_tokens=2000,
            temperature=0
        )
        
        # Parse response
        raw = response.choices[0].message.content
        
        if raw is None:
            raw = getattr(response.choices[0].message, "parsed", None)
        
        if isinstance(raw, list):
            content = "".join(
                (p.get("text", "") if isinstance(p, dict) else str(p)) for p in raw
            )
        elif raw is None:
            content = ""
        else:
            content = str(raw)
        
        # Extract JSON
        content = content.strip()
        
        try:
            explanation_data = json.loads(content)
        except:
            # Try to extract JSON from markdown code blocks
            m = re.search(r"```json\s*(\{.*?\})\s*```", content, re.S)
            if m:
                explanation_data = json.loads(m.group(1))
            else:
                m = re.search(r"(\{.*\})", content, re.S)
                if m:
                    explanation_data = json.loads(m.group(1))
                else:
                    raise ValueError("Could not parse JSON from response")
        
        # Normalize to expected schema
        normalized = {
            "titulo": explanation_data.get("titulo", f"Slide {slide_number}"),
            "explicacion_didactica": explanation_data.get("explicacion_didactica", ""),
            "puntos_clave": explanation_data.get("puntos_clave", []) or [],
            "conexiones": explanation_data.get("conexiones", ""),
            "resumen_corto": explanation_data.get("resumen_corto", ""),
            "anki_cards": explanation_data.get("anki_cards", []) or []
        }
        
        return {
            "success": True,
            "slide_number": slide_number,
            "explanation": normalized
        }
    
    except Exception as e:
        logger.error(f"Error analyzing slide {slide_number}: {e}")
        return {
            "success": False,
            "slide_number": slide_number,
            "error": str(e)
        }


def generate_summary_json(explanations: List[Dict]) -> Dict[str, Any]:
    """
    Generate a summary JSON from all slide explanations
    
    Args:
        explanations: List of explanation dicts
    
    Returns:
        Summary dictionary
    """
    summary = {
        "slides": [],
        "total_slides": len(explanations),
        "anki_cards": []
    }
    
    for exp in explanations:
        if exp.get("success"):
            slide_data = exp["explanation"]
            summary["slides"].append({
                "slide_number": exp["slide_number"],
                "titulo": slide_data.get("titulo", ""),
                "explicacion_didactica": slide_data.get("explicacion_didactica", ""),
                "puntos_clave": slide_data.get("puntos_clave", []),
                "conexiones": slide_data.get("conexiones", ""),
                "resumen_corto": slide_data.get("resumen_corto", "")
            })
            
            # Collect all anki cards
            anki_cards = slide_data.get("anki_cards", [])
            for card in anki_cards:
                summary["anki_cards"].append({
                    "slide_number": exp["slide_number"],
                    "front": card.get("pregunta", ""),
                    "back": card.get("respuesta", "")
                })
    
    return summary


def generate_docx(explanations: List[Dict], slides: List[bytes]) -> bytes:
    """
    Generate a Word document with slides and explanations
    
    Args:
        explanations: List of explanation dicts
        slides: List of slide image bytes
    
    Returns:
        DOCX file bytes
    """
    tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    tmp_docx.close()
    
    temp_images = []
    
    try:
        doc = Document()
        
        # Define styles
        title_style = doc.styles.add_style('SlideTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(0)
        
        heading_style = doc.styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        heading_style.paragraph_format.space_after = Pt(3)
        
        normal_style = doc.styles.add_style('NormalText', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.left_indent = Inches(0.25)
        normal_style.paragraph_format.space_after = Pt(3)
        
        # Process each slide
        for i, (slide_bytes, explanation) in enumerate(zip(slides, explanations)):
            slide_num = i + 1
            
            # Title
            title_para = doc.add_paragraph(f"Slide {slide_num}", style='SlideTitle')
            title_para.paragraph_format.space_before = Pt(6)
            
            # Add slide image
            try:
                img_tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                img_tmp.write(slide_bytes)
                img_tmp.close()
                temp_images.append(img_tmp.name)
                
                doc.add_picture(img_tmp.name, width=Inches(6))
            except Exception as e:
                doc.add_paragraph(f"Error loading slide image: {e}", style='NormalText')
            
            # Add explanation
            if explanation["success"]:
                exp_data = explanation["explanation"]
                
                titulo = exp_data.get('titulo', '')
                explicacion = exp_data.get('explicacion_didactica', '')
                puntos = exp_data.get('puntos_clave', [])
                conexiones = exp_data.get('conexiones', '')
                resumen_corto = exp_data.get('resumen_corto', '')
                
                if titulo:
                    title_para = doc.add_paragraph(style='SectionHeading')
                    title_para.add_run("📌 Título: ").bold = True
                    title_para.add_run(titulo).bold = True
                    title_para.paragraph_format.space_after = Pt(18)
                
                if explicacion:
                    doc.add_paragraph("🧠 Explicación didáctica", style='SectionHeading')
                    if isinstance(explicacion, list):
                        for item in explicacion:
                            doc.add_paragraph(item, style='NormalText')
                            doc.add_paragraph("", style='NormalText')
                    else:
                        doc.add_paragraph(explicacion, style='NormalText')
                    doc.add_paragraph("", style='NormalText')
                
                if puntos:
                    doc.add_paragraph("🎯 Puntos clave", style='SectionHeading')
                    for item in puntos:
                        para = doc.add_paragraph(f"• {item}", style='NormalText')
                        run = para.add_run()
                        run.add_break()
                
                if conexiones:
                    doc.add_paragraph("🔗 Conexiones", style='SectionHeading')
                    doc.add_paragraph(conexiones, style='NormalText')
                    doc.add_paragraph("", style='NormalText')
                
                if resumen_corto:
                    doc.add_paragraph("📝 Resumen corto", style='SectionHeading')
                    doc.add_paragraph(resumen_corto, style='NormalText')
                    doc.add_paragraph("", style='NormalText')
            else:
                doc.add_paragraph("❌ Error en el análisis", style='SectionHeading')
                doc.add_paragraph(explanation.get('error', 'Error desconocido'), style='NormalText')
            
            # Spacing between slides
            doc.add_paragraph("", style='NormalText')
            doc.add_paragraph("", style='NormalText')
            doc.add_paragraph("", style='NormalText')
        
        # Save document
        doc.save(tmp_docx.name)
        
        # Read bytes
        with open(tmp_docx.name, 'rb') as f:
            docx_bytes = f.read()
        
        return docx_bytes
    
    finally:
        # Cleanup
        try:
            os.unlink(tmp_docx.name)
        except:
            pass
        
        for img_file in temp_images:
            try:
                os.unlink(img_file)
            except:
                pass


def generate_anki_package(explanations: List[Dict]) -> bytes:
    """
    Generate Anki package (.apkg) from explanations
    
    Args:
        explanations: List of explanation dicts
    
    Returns:
        .apkg file bytes
    """
    # Create Anki model
    anki_model = genanki.Model(
        1607392319,
        'PDF Slide Explainer Model',
        fields=[
            {'name': 'Question'},
            {'name': 'Answer'},
            {'name': 'Source'},
        ],
        templates=[
            {
                'name': 'Card 1',
                'qfmt': '{{Question}}',
                'afmt': '{{FrontSide}}<hr id="answer">{{Answer}}<br><br><small style="color: #666;">{{Source}}</small>',
            },
        ],
        css="""
        .card {
            font-family: arial;
            font-size: 20px;
            text-align: center;
            color: black;
            background-color: white;
        }
        """
    )
    
    # Create deck
    anki_deck = genanki.Deck(2059400110, "Lecture Notes")
    
    # Add cards
    for explanation in explanations:
        if explanation.get("success"):
            exp_data = explanation["explanation"]
            anki_cards = exp_data.get("anki_cards", [])
            slide_num = explanation["slide_number"]
            titulo = exp_data.get('titulo', f'Slide {slide_num}')
            
            for card in anki_cards:
                if isinstance(card, dict):
                    pregunta = card.get('pregunta', '').strip()
                    respuesta = card.get('respuesta', '').strip()
                    
                    if pregunta and respuesta:
                        note = genanki.Note(
                            model=anki_model,
                            fields=[pregunta, respuesta, f"Slide {slide_num}: {titulo}"]
                        )
                        anki_deck.add_note(note)
    
    # Generate package
    package = genanki.Package(anki_deck)
    
    tmp_file = tempfile.NamedTemporaryFile(suffix='.apkg', delete=False)
    package.write_to_file(tmp_file.name)
    
    with open(tmp_file.name, 'rb') as f:
        apkg_bytes = f.read()
    
    try:
        os.unlink(tmp_file.name)
    except:
        pass
    
    return apkg_bytes


async def process_lecture(job_id: str, s3_key: str, email: str, language: str) -> Dict[str, Any]:
    """
    Main pipeline: download PDF, process, upload outputs, return URLs
    
    Args:
        job_id: Job ID
        s3_key: S3 key of the PDF
        email: User email
        language: Language for explanations
    
    Returns:
        Dictionary with output URLs
    """
    logger.info(f"[jobId={job_id}] Starting pipeline", extra={"jobId": job_id})
    
    # 1. Download PDF from S3
    logger.info(f"[jobId={job_id}] Downloading PDF from S3", extra={"jobId": job_id})
    pdf_bytes = download_from_s3(s3_key)
    
    # 2. Extract slides
    logger.info(f"[jobId={job_id}] Extracting slides", extra={"jobId": job_id})
    slides = extract_slides_from_pdf(pdf_bytes)
    
    if not slides:
        raise Exception("No slides extracted from PDF")
    
    # 3. Process each slide
    logger.info(f"[jobId={job_id}] Processing {len(slides)} slides", extra={"jobId": job_id})
    explanations = []
    
    for i, slide_bytes in enumerate(slides):
        slide_num = i + 1
        logger.info(f"[jobId={job_id}] Processing slide {slide_num}/{len(slides)}", extra={"jobId": job_id})
        explanation = explain_slide(slide_bytes, slide_num, language)
        explanations.append(explanation)
    
    # 4. Generate outputs
    logger.info(f"[jobId={job_id}] Generating output files", extra={"jobId": job_id})
    
    # JSON summary
    summary_json = generate_summary_json(explanations)
    summary_json_bytes = json.dumps(summary_json, indent=2, ensure_ascii=False).encode('utf-8')
    
    # DOCX
    docx_bytes = generate_docx(explanations, slides)
    
    # Anki package
    anki_bytes = generate_anki_package(explanations)
    
    # 5. Upload to S3
    logger.info(f"[jobId={job_id}] Uploading outputs to S3", extra={"jobId": job_id})
    
    summary_key = f"outputs/{job_id}/summary.json"
    docx_key = f"outputs/{job_id}/lecture.docx"
    anki_key = f"outputs/{job_id}/lecture.apkg"
    
    upload_to_s3(summary_json_bytes, summary_key, "application/json")
    upload_to_s3(docx_bytes, docx_key, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    upload_to_s3(anki_bytes, anki_key, "application/octet-stream")
    
    # 6. Generate presigned URLs
    logger.info(f"[jobId={job_id}] Generating presigned URLs", extra={"jobId": job_id})
    
    summary_url = generate_presigned_url(summary_key, expiration=86400)  # 24 hours
    docx_url = generate_presigned_url(docx_key, expiration=86400)
    anki_url = generate_presigned_url(anki_key, expiration=86400)
    
    logger.info(f"[jobId={job_id}] Pipeline complete", extra={"jobId": job_id})
    
    return {
        "summary_json_url": summary_url,
        "docx_url": docx_url,
        "anki_url": anki_url,
        "total_slides": len(slides)
    }

