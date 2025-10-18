"""PDF Slide Explainer using Flora Facturacion Infrastructure
Streamlit app that processes PDF slides and generates explanations for each slide
"""

import streamlit as st
import os
import base64
import json
from typing import List, Dict, Any, Optional
from openai import OpenAI
import fitz  # PyMuPDF for PDF processing
from PIL import Image
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import tempfile
import re
import genanki #libreria para generar ankis
import random
import difflib


def get_prompt(language: str = "Spanish") -> str:
    """Get the prompt template adapted for the specified language"""
    language_instruction = f"\n- Esta explicación debe ser escrita en {language}.\n"
    
    return f"""\
Hazme una explicación **completa, clara y dinámica** sobre este texto.{language_instruction}
Debe permitir al lector **entender todo el contenido técnico de manera fácil y ordenada**,
sin extenderse demasiado ni omitir ningún detalle importante.
Incluye ejemplos o analogías cuando ayuden a comprender mejor.
Al final, agrega un **resumen corto** con lo más importante de toda la explicación.

OBJETIVO GENERAL
- Que sea **profunda pero comprensible**, con rigor técnico y tono didáctico.
- Que ayude a **aprender de forma rápida**.
- Que combine explicación fluida con **puntos clave**.
- **IMPORTANTE:** Siempre genera texto explicativo completo, incluso si la diapositiva es un gráfico, diagrama o imagen sin texto. Analiza visualmente y describe lo que ves, explicando su significado y relevancia.
- **IDIOMA:** Todas las explicaciones deben estar en **{language}**, pero mantén las **palabras técnicas más importantes** (términos clave, conceptos específicos) en su **idioma original** (inglés, alemán, etc.) para facilitar el aprendizaje de vocabulario técnico.

INSTRUCCIONES
1) Explica el tema principal y por qué es relevante.
2) **EXPLICACIÓN DIDÁCTICA:** Divide la explicación completa en **puntos clave detallados y profundos** (no un párrafo largo). Cada punto debe ser **súper completo, técnico y profesional**, cubriendo TODOS los detalles visibles en la diapositiva sin omitir absolutamente nada. Explica conceptos complejos de manera que un principiante pueda entenderlos desde cero, pero con rigor técnico suficiente para convertir al lector en un experto absoluto que domine los conceptos y pueda usar términos técnicos correctamente. Incluye definiciones, ejemplos prácticos, analogías cuando ayuden, y conexiones lógicas. Mantén términos técnicos importantes en inglés o alemán si aplica, explicándolos en {language} cuando sea necesario. Usa más términos en inglés para conceptos clave y nombres específicos del PowerPoint, explicándolos en {language} cuando sea necesario. Proporciona el contenido directo sin prefijos como "Punto 1:", "Punto 2:", etc.
3) Resume conceptos principales adicionales en puntos clave (usando términos originales donde sea clave).
4) Conecta con temas relacionados, pero haciéndolo específico y en relación con las demás diapositivas, no tan general. Aporta información realmente útil y que ayude a comprender mejor el tema, no datos innecesarios.
5) Cierra con un **resumen corto** (2–3 frases con el takeaway).
6) **SIEMPRE** proporciona contenido completo, no dejes campos vacíos o con "N/A".
7) Genera 3 ankis de conceptos importantes de aprender e interiorizar en esta diapositiva.

FORMATO DE SALIDA
Devuelve **únicamente** un **objeto JSON válido** (sin texto adicional, sin comentarios).
NO copies literalmente el ejemplo; rellénalo con el contenido del slide.

```json
{{
  "titulo": "Tema o concepto central de la diapositiva",
  "explicacion_didactica": ["Punto 1: Detalle completo...", "Punto 2: Detalle completo...", "Punto 3: ..."],
  "puntos_clave": ["Idea 1", "Idea 2", "Idea 3"],
  "conexiones": "Relaciones con otros temas importantes",
  "resumen_corto": "Síntesis breve (2–3 frases)",
  "anki_cards": [
    {{"pregunta": "...", "respuesta": "..."}},
    {{"pregunta": "...", "respuesta": "..."}},
    {{"pregunta": "...", "respuesta": "..."}}
  ]
}}
"""


def encode_image_base64(image_bytes: bytes) -> str:
    """Encode image bytes to base64 string"""
    return base64.b64encode(image_bytes).decode('utf-8')

def init_openai_client(api_key: Optional[str] = None):
    """Initialize OpenAI client with API key"""
    if not api_key:
        api_key = os.getenv("OPENAI_API_KEY")
    
    if not api_key:
        return None
    
    try:
        return OpenAI(api_key=api_key)
    except Exception as e:
        st.error(f"❌ Error initializing OpenAI client: {str(e)}")
        return None

def extract_slides_from_pdf(pdf_file) -> List[bytes]:
    """
    Extract individual slides/pages from PDF as images
    
    Args:
        pdf_file: Uploaded PDF file from Streamlit
        
    Returns:
        List of image bytes for each slide
    """
    slides = []
    
    try:
        # Open PDF from uploaded file
        pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
        
        for page_num in range(len(pdf_document)):
            # Get page
            page = pdf_document.load_page(page_num)
            
            # Convert to image (300 DPI for good quality)
            mat = fitz.Matrix(300/72, 300/72)  # 300 DPI scaling
            pix = page.get_pixmap(matrix=mat)
            
            # Convert to PNG bytes
            img_data = pix.tobytes("png")
            slides.append(img_data)
            
        pdf_document.close()
        return slides
        
    except Exception as e:
        st.error(f"Error extracting slides from PDF: {str(e)}")
        return []

def explain_slide(slide_image_bytes: bytes, openai_client: OpenAI, slide_number: int, custom_prompt: Optional[str] = None, selected_language: str = "Spanish") -> Dict[str, Any]:
    """
    Generate explanation for a single slide using OpenAI Vision API
    
    Args:
        slide_image_bytes: Image bytes of the slide
        openai_client: OpenAI client instance
        slide_number: Number of the slide (for context)
        
    Returns:
        Dictionary with slide explanation and analysis
    """
    try:
        # Encode image to base64
        image_base64 = encode_image_base64(slide_image_bytes)
        image_url = f"data:image/png;base64,{image_base64}"
        
        # Use custom prompt if provided, otherwise use default with language adaptation
        # IMPORTANT: Avoid str.format here because prompt templates contain JSON braces
        # which would be interpreted as format fields. We only want to substitute {slide_number}.
        if custom_prompt:
            explanation_prompt = custom_prompt.replace("{slide_number}", str(slide_number))
        else:
            explanation_prompt = get_prompt(selected_language)

        
        # Call Vision API
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            response_format={"type": "json_object"},  # <— NUEVO: fuerza JSON puro
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": explanation_prompt},
                        {"type": "image_url", "image_url": {"url": image_url, "detail": "high"}}
                    ]
                }
            ],
            max_tokens=2000,
            temperature=0  # <— recomendado para consistencia
        )
        
        # Parse response (normaliza a string y extrae JSON de forma robusta)
        raw = response.choices[0].message.content

        # Algunos SDK/dev builds pueden devolver None o listas de partes
        if raw is None:
            # intenta recuperar de un posible atributo alternativo
            raw = getattr(response.choices[0].message, "parsed", None)

        if isinstance(raw, list):
            content = "".join(
                (p.get("text", "") if isinstance(p, dict) else str(p)) for p in raw
            )
        elif raw is None:
            content = ""
        else:
            content = str(raw)

        def extract_json_safe(s: str):
            # Clean input first
            s = s.strip()
            
            # 1) JSON puro
            try:
                return json.loads(s)
            except Exception:
                pass
                
            # 2) Try to fix malformed JSON that starts with quotes
            if s.startswith('"') and not s.startswith('{"'):
                try:
                    # Try adding opening brace
                    fixed = "{" + s
                    return json.loads(fixed)
                except Exception:
                    try:
                        # Try removing leading quotes and finding JSON-like content
                        cleaned = s.lstrip('\n "')
                        if ':' in cleaned:
                            fixed = '{"' + cleaned
                            return json.loads(fixed)
                    except Exception:
                        pass
                    
            # 3) bloque ```json ... ```
            m = re.search(r"```json\s*(\{.*?\})\s*```", s, re.S)
            if m:
                try:
                    return json.loads(m.group(1))
                except Exception:
                    pass
                    
            # 4) bloque ``` ... ```
            m = re.search(r"```\s*(\{.*?\})\s*```", s, re.S)
            if m:
                try:
                    return json.loads(m.group(1))
                except Exception:
                    pass
                    
            # 5) primer objeto { ... }
            m = re.search(r"(\{.*\})", s, re.S)
            if m:
                try:
                    return json.loads(m.group(1))
                except Exception:
                    pass
                    
            # 6) último recurso: limpia ecos de {{ ... }} y usa como explicación
            cleaned = re.sub(r"\{\{[\s\S]*?\}\}", "", s).strip()
            return {
                "titulo": f"Slide {slide_number}",
                "explicacion_didactica": cleaned if cleaned else s,
                "puntos_clave": [],
                "conexiones": "",
                "resumen_corto": ""
            }

        explanation_data = extract_json_safe(content)

        
        # === Normalización de esquema al nuevo formato ===
        # Si ya viene en el esquema nuevo, lo usamos tal cual:
        if all(k in explanation_data for k in ["titulo", "explicacion_didactica", "puntos_clave", "conexiones", "resumen_corto"]):
            normalized = {
                "titulo": explanation_data.get("titulo", ""),
                "explicacion_didactica": explanation_data.get("explicacion_didactica", ""),
                "puntos_clave": explanation_data.get("puntos_clave", []) or [],
                "conexiones": explanation_data.get("conexiones", ""),
                "resumen_corto": explanation_data.get("resumen_corto", ""),
                "anki_cards": explanation_data.get("anki_cards", []) or []
            }
        else:
            # Fallback desde el esquema antiguo
            titulo_old = explanation_data.get("titulo", f"Slide {slide_number}")
            contenido_clave_old = explanation_data.get("contenido_clave", [])
            contexto_old = explanation_data.get("contexto", "")
            insights_old = explanation_data.get("insights", [])
            resumen_old = explanation_data.get("resumen", "")

            # Construimos la explicación didáctica a partir de lo disponible
            explicacion_didactica_new = ""
            if isinstance(resumen_old, str) and resumen_old.strip():
                explicacion_didactica_new = resumen_old.strip()
            else:
                parts = []
                if isinstance(contenido_clave_old, list) and contenido_clave_old:
                    parts.append(" ".join(contenido_clave_old))
                if isinstance(contexto_old, str) and contexto_old.strip():
                    parts.append(contexto_old.strip())
                if isinstance(insights_old, list) and insights_old:
                    parts.append(" ".join(insights_old))
                explicacion_didactica_new = " ".join(p for p in parts if p).strip()

            normalized = {
                "titulo": titulo_old,
                "explicacion_didactica": explicacion_didactica_new or "Explicación generada automáticamente.",
                "puntos_clave": contenido_clave_old if isinstance(contenido_clave_old, list) else [],
                "conexiones": contexto_old if isinstance(contexto_old, str) else "",
                "resumen_corto": resumen_old if isinstance(resumen_old, str) else "",
                "anki_cards": explanation_data.get("anki_cards", []) or []
            }

        return {
            "success": True,
            "slide_number": slide_number,
            "explanation": normalized,
            "raw_response": content
        }

        
    except Exception as e:
        return {
            "success": False,
            "slide_number": slide_number,
            "error": f"Error analyzing slide {slide_number}: {str(e)}"
        }

def generate_word_report(slides: List[bytes], explanations: List[Dict], pdf_name: Optional[str]) -> bytes:
    """
    Generate a Word document (.docx) report with slides and explanations in continuous format

    Args:
        slides: List of slide image bytes
        explanations: List of explanation dictionaries
        pdf_name: Original PDF name for the report

    Returns:
        Word document bytes
    """
    if not slides or not explanations:
        raise ValueError("Slides and explanations cannot be None or empty")

    # Create temporary file for Word document
    tmp_docx_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    tmp_docx_file.close()

    # Keep track of temporary image files to clean up later
    temp_image_files = []

    try:
        # Create Word document
        doc = Document()

        # Set up styles
        title_style = doc.styles.add_style('SlideTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(18)  # type: ignore
        title_style.font.bold = True  # type: ignore
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # type: ignore
        title_style.paragraph_format.space_after = Pt(0)  # type: ignore

        heading_style = doc.styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.size = Pt(14)  # type: ignore
        heading_style.font.bold = True  # type: ignore
        heading_style.paragraph_format.space_after = Pt(3)  # type: ignore

        normal_style = doc.styles.add_style('NormalText', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.size = Pt(11)  # type: ignore
        normal_style.paragraph_format.left_indent = Inches(0.25)  # type: ignore
        normal_style.paragraph_format.space_after = Pt(3)  # type: ignore

        # Process each slide
        for i, (slide_bytes, explanation) in enumerate(zip(slides, explanations)):
            slide_num = i + 1

            # Slide title with minimal space before
            title_para = doc.add_paragraph(f"Slide {slide_num}", style='SlideTitle')
            title_para.paragraph_format.space_before = Pt(6)  # Reduced space before slide title

            # Add slide image
            try:
                # Create temporary image file
                img_tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                img_tmp.write(slide_bytes)
                img_tmp.close()

                temp_image_files.append(img_tmp.name)  # Track for cleanup

                # Add image to document (width: 6 inches, height: auto-maintain aspect ratio)
                doc.add_picture(img_tmp.name, width=Inches(6))

                # No extra space after image - keep content close to slide

            except Exception as e:
                error_para = doc.add_paragraph(f"Error loading slide image: {str(e)}", style='NormalText')

            # Add explanation
            if explanation["success"]:
                exp_data = explanation["explanation"]

                title = exp_data.get('titulo') or ""
                exp_did = exp_data.get('explicacion_didactica') or ""
                resumen = exp_data.get('resumen') or ""
                resumen_corto = exp_data.get('resumen_corto') or ""
                puntos = exp_data.get('puntos_clave') or exp_data.get('contenido_clave') or []
                conex = exp_data.get('conexiones') or exp_data.get('contexto') or ""

                explicacion = exp_did if (isinstance(exp_did, str) and exp_did.strip()) or (isinstance(exp_did, list) and exp_did) else resumen

                if title:
                    # Create a paragraph with "📌 Título:" and the title in the same line, like the web interface
                    title_para = doc.add_paragraph(style='SectionHeading')
                    title_para.add_run("📌 Título: ").bold = True
                    title_para.add_run(title).bold = True
                    title_para.paragraph_format.space_after = Pt(18)  # Add more space after

                if explicacion:
                    doc.add_paragraph("🧠 Explicación didáctica", style='SectionHeading')
                    if isinstance(explicacion, list):
                        for item in explicacion:
                            doc.add_paragraph(item, style='NormalText')
                            doc.add_paragraph("", style='NormalText')  # Salto de línea entre puntos
                    else:
                        doc.add_paragraph(explicacion, style='NormalText')
                    doc.add_paragraph("", style='NormalText')  # Salto de línea

                if puntos:
                    doc.add_paragraph("🎯 Puntos clave", style='SectionHeading')
                    for item in puntos:
                        para = doc.add_paragraph(f"• {item}", style='NormalText')
                        run = para.add_run()
                        run.add_break()  # Salto de línea delicado entre puntos clave

                if conex:
                    doc.add_paragraph("🔗 Conexiones", style='SectionHeading')
                    doc.add_paragraph(conex, style='NormalText')
                    doc.add_paragraph("", style='NormalText')  # Salto de línea

                # Solo muestra 'Resumen' si es distinto de la explicación
                if resumen and isinstance(resumen, str) and resumen.strip() != (explicacion.strip() if isinstance(explicacion, str) else ""):
                    doc.add_paragraph("📝 Resumen", style='SectionHeading')
                    doc.add_paragraph(resumen, style='NormalText')
                    doc.add_paragraph("", style='NormalText')  # Salto de línea

                # Solo muestra 'Resumen corto' si es distinto
                if resumen_corto and isinstance(resumen_corto, str) and resumen_corto.strip() not in {(resumen.strip() if isinstance(resumen, str) else ""), (explicacion.strip() if isinstance(explicacion, str) else "")}:
                    doc.add_paragraph("📝 Resumen corto", style='SectionHeading')
                    doc.add_paragraph(resumen_corto, style='NormalText')
                    doc.add_paragraph("", style='NormalText')  # Salto de línea

            else:
                doc.add_paragraph("❌ Error en el análisis", style='SectionHeading')
                doc.add_paragraph(explanation.get('error', 'Error desconocido'), style='NormalText')

            # Add space between slides instead of page break for better copy-paste compatibility
            doc.add_paragraph("", style='NormalText')
            doc.add_paragraph("", style='NormalText')
            doc.add_paragraph("", style='NormalText')  # Three empty paragraphs for clear separation

        # Save the document
        doc.save(tmp_docx_file.name)

        # Read the generated Word document
        with open(tmp_docx_file.name, 'rb') as f:
            docx_bytes = f.read()

        return docx_bytes

    finally:
        # Clean up temporary files
        try:
            os.unlink(tmp_docx_file.name)
        except:
            pass

        for img_file in temp_image_files:
            try:
                os.unlink(img_file)
            except:
                pass

def generate_anki_export(explanations: List[Dict], pdf_name: Optional[str]) -> bytes:
    """
    Generate Anki deck (.apkg) file from slide explanations using genanki

    Args:
        explanations: List of explanation dictionaries containing anki_cards
        pdf_name: Original PDF name for context

    Returns:
        Bytes of the .apkg file
    """
    if not explanations:
        return b""

    # Create Anki model (card template)
    anki_model = genanki.Model(
        1607392319,  # Hardcoded unique model ID
        'PDF Slide Explainer Model',
        fields=[
            {'name': 'Question'},
            {'name': 'Answer'},
            {'name': 'Source'},
        ],
        templates=[
            {
                'name': 'Card 1',
                'qfmt': '{{Question}}',
                'afmt': '{{FrontSide}}<hr id="answer">{{Answer}}<br><br><small style="color: #666;">{{Source}}</small>',
            },
        ],
        css="""
        .card {
            font-family: arial;
            font-size: 20px;
            text-align: center;
            color: black;
            background-color: white;
        }
        .cloze {
            font-weight: bold;
            color: blue;
        }
        """
    )

    # Create deck
    deck_name = f"PDF Slide Explainer - {pdf_name}" if pdf_name else "PDF Slide Explainer"
    anki_deck = genanki.Deck(
        2059400110,  # Hardcoded unique deck ID
        deck_name 
    )

    # Add notes to deck
    card_count = 0
    for i, explanation in enumerate(explanations):
        if explanation.get("success") and explanation.get("explanation"):
            exp_data = explanation["explanation"]
            anki_cards = exp_data.get("anki_cards", [])
            slide_title = exp_data.get('titulo', f'Slide {i + 1}')

            if anki_cards:
                for card in anki_cards:
                    if isinstance(card, dict) and 'pregunta' in card and 'respuesta' in card:
                        # Create Anki note
                        note = genanki.Note(
                            model=anki_model,
                            fields=[
                                card['pregunta'].strip(),
                                card['respuesta'].strip(),
                                f"Slide {i + 1}: {slide_title}"
                            ]
                        )
                        anki_deck.add_note(note)
                        card_count += 1

    # Generate .apkg file in memory
    package = genanki.Package(anki_deck)

    # Create temporary file to get bytes
    with tempfile.NamedTemporaryFile(suffix='.apkg', delete=False) as tmp_file:
        package.write_to_file(tmp_file.name)

        # Read the file content
        with open(tmp_file.name, 'rb') as f:
            apkg_bytes = f.read()

        # Clean up temp file
        try:
            os.unlink(tmp_file.name)
        except:
            pass

    return apkg_bytes

def generate_quiz(anki_cards_list: List[Dict]) -> List[Dict]:
    """
    Generate a quiz with 20 multiple choice questions from Anki cards

    Args:
        anki_cards_list: List of all anki cards from all slides

    Returns:
        List of quiz questions, each with question, options, correct_answer
    """
    if not anki_cards_list:
        return []

    # Filter valid cards
    valid_cards = [card for card in anki_cards_list if isinstance(card, dict) and 'pregunta' in card and 'respuesta' in card]

    if len(valid_cards) < 4:  # Need at least 4 for multiple choice
        return []

    # Select up to 20 questions randomly
    num_questions = min(20, len(valid_cards))
    selected_cards = random.sample(valid_cards, num_questions)

    quiz_questions = []

    for card in selected_cards:
        question = card['pregunta'].strip()
        correct_answer = card['respuesta'].strip()

        # Get all other answers for distractors
        other_answers = [c['respuesta'].strip() for c in valid_cards if c != card and c['respuesta'].strip() != correct_answer]

        if len(other_answers) < 3:
            continue  # Skip if not enough distractors

        # Find 3 most similar answers using difflib
        similarities = [(ans, difflib.SequenceMatcher(None, correct_answer, ans).ratio()) for ans in other_answers]
        similarities.sort(key=lambda x: x[1], reverse=True)
        distractors = [sim[0] for sim in similarities[:3]]

        # Create options: correct + 3 distractors, shuffled
        options = [correct_answer] + distractors
        random.shuffle(options)

        quiz_questions.append({
            'question': question,
            'options': options,
            'correct_answer': correct_answer
        })

    return quiz_questions

def main():
    """Main Streamlit application"""
    st.set_page_config(
        page_title="PDF Slide Explainer",
        page_icon="📊",
        layout="wide"
    )

    # Custom CSS for modern appearance
    st.markdown("""
    <style>
    .main {
        background: linear-gradient(135deg, #A8DADC 0%, #7FB3D5 25%, #4A90E2 50%, #357ABD 75%, #1E3A8A 100%);
        color: #212529;
        min-height: 100vh;
    }
    .stTitle {
        color: #ffffff;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        font-weight: 700;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
    }
    .stMarkdown {
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        color: #ffffff;
    }
    .stButton>button {
        background: linear-gradient(135deg, #FF6B6B 0%, #FF8E53 25%, #FF6B35 50%, #4ECDC4 75%, #26D0CE 100%);
        color: white;
        border: none;
        border-radius: 30px;
        padding: 16px 32px;
        font-size: 16px;
        font-weight: 800;
        transition: all 0.4s ease;
        box-shadow: 0 8px 25px rgba(0,0,0,0.4), inset 0 1px 0 rgba(255,255,255,0.2);
        text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
        border: 1px solid rgba(255,255,255,0.1);
    }
    .stButton>button:hover {
        transform: translateY(-4px) scale(1.02);
        box-shadow: 0 12px 35px rgba(0,0,0,0.5), inset 0 1px 0 rgba(255,255,255,0.3);
        background: linear-gradient(135deg, #FF8E53 0%, #FF6B35 25%, #FF6B6B 50%, #4ECDC4 75%, #26D0CE 100%);
    }
    .stExpander {
        background: linear-gradient(135deg, rgba(10,10,25,0.98), rgba(20,20,45,0.96), rgba(35,35,65,0.94), rgba(50,50,85,0.92), rgba(65,65,105,0.9), rgba(80,80,125,0.88));
        border: 2px solid rgba(255, 255, 255, 0.4);
        border-radius: 20px;
        box-shadow: 0 15px 50px rgba(0,0,0,0.5), inset 0 1px 0 rgba(255,255,255,0.1);
        backdrop-filter: blur(20px);
        margin-bottom: 25px;
    }
    .stExpanderHeader {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 25%, #4A90E2 50%, #357ABD 75%, #1E3A8A 100%);
        color: white;
        border-radius: 18px 18px 0 0;
        font-weight: 800;
        padding: 20px;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.4);
        border-bottom: 1px solid rgba(255,255,255,0.2);
    }
    .stTextInput>div>div>input {
        border-radius: 10px;
        border: 2px solid rgba(255, 255, 255, 0.3);
        background: rgba(255, 255, 255, 0.9);
        color: #333;
        padding: 10px;
    }
    .stTextArea>div>div>textarea {
        border-radius: 10px;
        border: 2px solid rgba(255, 255, 255, 0.3);
        background: rgba(255, 255, 255, 0.9);
        color: #333;
        padding: 10px;
    }
    .stSidebar {
        background: linear-gradient(135deg, #1a1a2e 0%, #16213e 25%, #0f0f23 50%, #0a0a1a 75%, #050510 100%);
        border-right: 4px solid rgba(74, 144, 226, 0.4);
        padding: 25px;
        box-shadow: inset 0 0 60px rgba(0,0,0,0.6), 0 0 30px rgba(74, 144, 226, 0.1);
    }
    .stSidebar .stMarkdown {
        color: #ecf0f1;
        font-weight: 500;
    }
    .stSidebar .stButton>button {
        background: linear-gradient(135deg, #3498db 0%, #4A90E2 50%, #5DADE2 100%);
        margin-bottom: 12px;
        box-shadow: 0 4px 15px rgba(52, 152, 219, 0.3);
    }
    .stSidebar .stButton>button:hover {
        background: linear-gradient(135deg, #5DADE2 0%, #3498db 50%, #4A90E2 100%);
        box-shadow: 0 6px 20px rgba(52, 152, 219, 0.4);
    }
    .stSidebar .stTextInput label, .stSidebar .stTextArea label {
        color: #ecf0f1 !important;
        font-weight: 600;
    }
    .stImage {
        border-radius: 15px;
        box-shadow: 0 8px 25px rgba(0,0,0,0.3);
        border: 2px solid rgba(255,255,255,0.1);
    }
    .css-1d391kg {  /* Main container */
        background: transparent;
    }
    </style>
    """, unsafe_allow_html=True)

    # Enhanced title section with modern design
    st.markdown("""
    <div style="text-align: center; padding: 40px 20px; margin-bottom: 20px;">
        <div style="background: linear-gradient(135deg, rgba(15,15,35,0.98), rgba(25,25,50,0.95), rgba(40,40,70,0.92)); padding: 30px; border-radius: 20px; box-shadow: 0 15px 40px rgba(0,0,0,0.4); backdrop-filter: blur(15px); border: 1px solid rgba(255,255,255,0.1);">
            <h1 style="color: #ffffff; font-size: 3.5em; margin-bottom: 10px; text-shadow: 2px 2px 4px rgba(0,0,0,0.5); font-weight: 800;">
                📊 PDF Slide Explainer
            </h1>
            <div style="background: linear-gradient(45deg, rgba(255,255,255,0.1), rgba(255,255,255,0.05)); border: 2px solid rgba(255,255,255,0.3); display: inline-block; padding: 8px 18px; border-radius: 15px; margin: 15px 0; box-shadow: 0 2px 10px rgba(0,0,0,0.2);">
                <p style="color: #ffffff; font-size: 1em; margin: 0; font-weight: 600; letter-spacing: 0.3px;">
                    🚀 Powered by Flora Facturación Infrastructure
                </p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize session state
    if 'slides' not in st.session_state:
        st.session_state.slides = None
    if 'explanations' not in st.session_state:
        st.session_state.explanations = None
    if 'edited_explanations' not in st.session_state:
        st.session_state.edited_explanations = None
    if 'uploaded_file_name' not in st.session_state:
        st.session_state.uploaded_file_name = None
    if 'word_report' not in st.session_state:
        st.session_state.word_report = None
    if 'undo_stack' not in st.session_state:
        st.session_state.undo_stack = []
    if 'redo_stack' not in st.session_state:
        st.session_state.redo_stack = []
    if 'current_slide_view' not in st.session_state:
        st.session_state.current_slide_view = None
    if 'selected_language' not in st.session_state:
        st.session_state.selected_language = "Spanish"
    
    # Sidebar for configuration
    with st.sidebar:
        st.markdown("""
        <div style="text-align: center; padding: 15px 0; margin-bottom: 20px; border-bottom: 2px solid rgba(255,255,255,0.2);">
            <h2 style="color: #ffffff; margin: 0; font-size: 1.6em; font-weight: 700; text-shadow: 1px 1px 2px rgba(0,0,0,0.5); font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;">
                ⚙️ Configuration
            </h2>
        </div>
        """, unsafe_allow_html=True)

        # API Key section
        st.markdown("""
        <div style="margin-bottom: 20px;">
            <div style="background: linear-gradient(135deg, rgba(220,220,240,0.95), rgba(200,200,230,0.9), rgba(180,180,220,0.85)); padding: 20px; border-radius: 15px; box-shadow: 0 10px 30px rgba(0,0,0,0.2); backdrop-filter: blur(10px); border: 1px solid rgba(0,0,0,0.1);">
                <h4 style="color: #1a1a2e; font-size: 1.1em; margin: 0; text-shadow: 1px 1px 2px rgba(255,255,255,0.4); font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;">
                    🔑 OpenAI API Key
                </h4>
            </div>
        </div>
        """, unsafe_allow_html=True)
        api_key_input = st.text_input(
            "",
            type="password",
            help="Your API key will not be stored. Get one at https://platform.openai.com/api-keys",
            label_visibility="collapsed"
        )

        st.markdown("---")  # Separator

        # Custom prompt section with modern design
        st.markdown("""
        <div style="margin-bottom: 20px;">
            <div style="background: linear-gradient(135deg, rgba(220,220,240,0.95), rgba(200,200,230,0.9), rgba(180,180,220,0.85)); padding: 20px; border-radius: 15px; box-shadow: 0 10px 30px rgba(0,0,0,0.2); backdrop-filter: blur(10px); border: 1px solid rgba(0,0,0,0.1);">
                <h4 style="color: #1a1a2e; font-size: 1.1em; margin: 0 0 15px 0; text-shadow: 1px 1px 2px rgba(255,255,255,0.4); font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;">
                    🎯 Custom Analysis Prompt
                </h4>
            </div>
        </div>
        """, unsafe_allow_html=True)
        use_custom_prompt = st.checkbox("Use custom prompt", help="Customize the AI analysis prompt")

        custom_prompt = None
        if use_custom_prompt:
            custom_prompt = st.text_area(
                "Custom Prompt:",
                height=200,
                value=get_prompt("Spanish"), ## muestra en español por default por ahora, cuando salgamos a vender en ingles
                help="Use {slide_number} as placeholder for slide number",
                label_visibility="collapsed"
            )
    
    # Hero section with modern design (closer spacing)
    st.markdown("""
    <div style="text-align: center; padding: 30px 20px; background: linear-gradient(135deg, rgba(15,15,35,0.95), rgba(25,25,50,0.92), rgba(40,40,70,0.9), rgba(50,50,80,0.85)); border-radius: 20px; margin: 10px 0 30px 0; box-shadow: 0 20px 60px rgba(0,0,0,0.5), inset 0 1px 0 rgba(255,255,255,0.1); backdrop-filter: blur(20px); border: 1px solid rgba(255,255,255,0.15); position: relative; overflow: hidden;">
        <div style="position: absolute; top: 0; left: 0; right: 0; bottom: 0; background: radial-gradient(circle at 30% 20%, rgba(255,255,255,0.1) 0%, transparent 50%), radial-gradient(circle at 70% 80%, rgba(255,255,255,0.05) 0%, transparent 50%); pointer-events: none;"></div>
        <h2 style="color: #ffffff; font-size: 2.5em; margin-bottom: 10px; text-shadow: 2px 2px 4px rgba(0,0,0,0.7); font-weight: 800; position: relative; z-index: 1;">
            Transform your presentations into deep knowledge
        </h2>
        <p style="color: #ffffff; font-size: 1.2em; margin-bottom: 0; line-height: 1.6; position: relative; z-index: 1;">
            Upload your PDF and discover detailed explanations for each slide using advanced artificial intelligence
        </p>
    </div>
    """, unsafe_allow_html=True)

    # Features section with modern cards
    st.markdown("""
    <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)); gap: 20px; margin: 30px 0;">
        <div style="background: linear-gradient(135deg, rgba(15,15,35,0.95), rgba(25,25,50,0.92), rgba(40,40,70,0.9), rgba(50,50,80,0.85)); padding: 25px; border-radius: 15px; box-shadow: 0 15px 40px rgba(0,0,0,0.4), inset 0 1px 0 rgba(255,255,255,0.1); backdrop-filter: blur(15px); border: 1px solid rgba(255,255,255,0.15); position: relative; overflow: hidden;">
            <div style="position: absolute; top: 0; left: 0; right: 0; bottom: 0; background: radial-gradient(circle at 20% 30%, rgba(255,255,255,0.08) 0%, transparent 40%); pointer-events: none;"></div>
            <h3 style="color: #ffffff; margin-bottom: 15px; text-shadow: 1px 1px 2px rgba(0,0,0,0.6); position: relative; z-index: 1;">🧠 Intelligent Analysis</h3>
            <p style="color: #ffffff; line-height: 1.6; position: relative; z-index: 1;">We use GPT-4 Vision to analyze every visual and textual element of your slides, generating complete and pedagogical explanations.</p>
        </div>
        <div style="background: linear-gradient(135deg, rgba(15,15,35,0.95), rgba(25,25,50,0.92), rgba(40,40,70,0.9), rgba(50,50,80,0.85)); padding: 25px; border-radius: 15px; box-shadow: 0 15px 40px rgba(0,0,0,0.4), inset 0 1px 0 rgba(255,255,255,0.1); backdrop-filter: blur(15px); border: 1px solid rgba(255,255,255,0.15); position: relative; overflow: hidden;">
            <div style="position: absolute; top: 0; left: 0; right: 0; bottom: 0; background: radial-gradient(circle at 80% 20%, rgba(255,255,255,0.06) 0%, transparent 40%); pointer-events: none;"></div>
            <h3 style="color: #ffffff; margin-bottom: 15px; text-shadow: 1px 1px 2px rgba(0,0,0,0.6); position: relative; z-index: 1;">📋 Clear Structure</h3>
            <p style="color: #ffffff; line-height: 1.6; position: relative; z-index: 1;">Each explanation is organized in detailed key points, facilitating learning and deep understanding of complex concepts.</p>
        </div>
        <div style="background: linear-gradient(135deg, rgba(15,15,35,0.95), rgba(25,25,50,0.92), rgba(40,40,70,0.9), rgba(50,50,80,0.85)); padding: 25px; border-radius: 15px; box-shadow: 0 15px 40px rgba(0,0,0,0.4), inset 0 1px 0 rgba(255,255,255,0.1); backdrop-filter: blur(15px); border: 1px solid rgba(255,255,255,0.15); position: relative; overflow: hidden;">
            <div style="position: absolute; top: 0; left: 0; right: 0; bottom: 0; background: radial-gradient(circle at 30% 70%, rgba(255,255,255,0.07) 0%, transparent 40%); pointer-events: none;"></div>
            <h3 style="color: #ffffff; margin-bottom: 15px; text-shadow: 1px 1px 2px rgba(0,0,0,0.6); position: relative; z-index: 1;">🎨 Visual Recognition</h3>
            <p style="color: #ffffff; line-height: 1.6; position: relative; z-index: 1;">We identify and explain graphics, diagrams and images, connecting visual elements with their conceptual meaning.</p>
        </div>
        <div style="background: linear-gradient(135deg, rgba(15,15,35,0.95), rgba(25,25,50,0.92), rgba(40,40,70,0.9), rgba(50,50,80,0.85)); padding: 25px; border-radius: 15px; box-shadow: 0 15px 40px rgba(0,0,0,0.4), inset 0 1px 0 rgba(255,255,255,0.1); backdrop-filter: blur(15px); border: 1px solid rgba(255,255,255,0.15); position: relative; overflow: hidden;">
            <div style="position: absolute; top: 0; left: 0; right: 0; bottom: 0; background: radial-gradient(circle at 70% 80%, rgba(255,255,255,0.05) 0%, transparent 40%); pointer-events: none;"></div>
            <h3 style="color: #ffffff; margin-bottom: 15px; text-shadow: 1px 1px 2px rgba(0,0,0,0.6); position: relative; z-index: 1;">📊 Professional Reports</h3>
            <p style="color: #ffffff; line-height: 1.6; position: relative; z-index: 1;">Generate polished Word documents with all explanations, perfect for sharing or archiving.</p>
        </div>
        <div style="background: linear-gradient(135deg, rgba(15,15,35,0.95), rgba(25,25,50,0.92), rgba(40,40,70,0.9), rgba(50,50,80,0.85)); padding: 25px; border-radius: 15px; box-shadow: 0 15px 40px rgba(0,0,0,0.4), inset 0 1px 0 rgba(255,255,255,0.1); backdrop-filter: blur(15px); border: 1px solid rgba(255,255,255,0.15); position: relative; overflow: hidden;">
            <div style="position: absolute; top: 0; left: 0; right: 0; bottom: 0; background: radial-gradient(circle at 40% 20%, rgba(255,255,255,0.06) 0%, transparent 40%); pointer-events: none;"></div>
            <h3 style="color: #ffffff; margin-bottom: 15px; text-shadow: 1px 1px 2px rgba(0,0,0,0.6); position: relative; z-index: 1;">✏️ Interactive Editing</h3>
            <p style="color: #ffffff; line-height: 1.6; position: relative; z-index: 1;">Customize generated explanations with inline editing tools and undo/redo functions.</p>
        </div>
        <div style="background: linear-gradient(135deg, rgba(15,15,35,0.95), rgba(25,25,50,0.92), rgba(40,40,70,0.9), rgba(50,50,80,0.85)); padding: 25px; border-radius: 15px; box-shadow: 0 15px 40px rgba(0,0,0,0.4), inset 0 1px 0 rgba(255,255,255,0.1); backdrop-filter: blur(15px); border: 1px solid rgba(255,255,255,0.15); position: relative; overflow: hidden;">
            <div style="position: absolute; top: 0; left: 0; right: 0; bottom: 0; background: radial-gradient(circle at 60% 60%, rgba(255,255,255,0.07) 0%, transparent 40%); pointer-events: none;"></div>
            <h3 style="color: #ffffff; margin-bottom: 15px; text-shadow: 1px 1px 2px rgba(0,0,0,0.6); position: relative; z-index: 1;">🔍 Detailed Exploration</h3>
            <p style="color: #ffffff; line-height: 1.6; position: relative; z-index: 1;">Enlarge slides to see details, remove unnecessary content and navigate efficiently through your analyses.</p>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize OpenAI client
    openai_client = init_openai_client(api_key_input)
    
    if not openai_client:
        if not api_key_input:
            st.warning("⚠️ Please enter your OpenAI API Key in the sidebar to continue.")
        else:
            st.error("❌ Invalid OpenAI API Key. Please check your key.")
        st.stop()
    else:
        st.success("✅ OpenAI client initialized successfully")
    
    # File upload
    st.markdown("**📁 Upload PDF Presentation**")  # Changed from subheader to markdown for less space
    uploaded_file = st.file_uploader(
        "Choose a PDF file",
        type=['pdf'],
        help="Upload a PDF presentation to analyze each slide"
    )
    
    if uploaded_file is not None:
        st.success(f"✅ File uploaded: {uploaded_file.name}")
        
        # Check if this is a new file or same file
        if st.session_state.uploaded_file_name != uploaded_file.name:
            # New file - clear previous results
            st.session_state.slides = None
            st.session_state.explanations = None
            st.session_state.word_report = None
            st.session_state.uploaded_file_name = uploaded_file.name
            
            # Extract slides
            with st.spinner("🔄 Extracting slides from PDF..."):
                st.session_state.slides = extract_slides_from_pdf(uploaded_file)
        
        if not st.session_state.slides:
            st.error("❌ Failed to extract slides from PDF")
            return
        
        st.success(f"✅ Extracted {len(st.session_state.slides)} slides")
        
        # Process slides
        if st.session_state.explanations is None:
            # Language selection
            st.markdown("**🌐 Language Selection**")
            language_options = {
                "Spanish": "🇪🇸 Spanish",
                "English": "🇺🇸 English", 
                "French": "🇫🇷 French",
                "German": "🇩🇪 German",
                "Italian": "🇮🇹 Italian",
                "Portuguese": "🇵🇹 Portuguese",
                "Chinese": "🇨🇳 Chinese",
                "Japanese": "🇯🇵 Japanese",
                "Korean": "🇰🇷 Korean",
                "Arabic": "🇸🇦 Arabic",
                "Russian": "🇷🇺 Russian",
                "Dutch": "🇳🇱 Dutch",
                "Swedish": "🇸🇪 Swedish",
                "Norwegian": "🇳🇴 Norwegian",
                "Danish": "🇩🇰 Danish",
                "Finnish": "🇫🇮 Finnish",
                "Polish": "🇵🇱 Polish",
                "Czech": "🇨🇿 Czech",
                "Hungarian": "🇭🇺 Hungarian",
                "Romanian": "🇷🇴 Romanian",
                "Greek": "🇬🇷 Greek",
                "Turkish": "🇹🇷 Turkish",
                "Hebrew": "🇮🇱 Hebrew",
                "Hindi": "🇮🇳 Hindi",
                "Thai": "🇹🇭 Thai",
                "Vietnamese": "🇻🇳 Vietnamese",
                "Indonesian": "🇮🇩 Indonesian",
                "Malay": "🇲🇾 Malay",
                "Filipino": "🇵🇭 Filipino",
                "Ukrainian": "🇺🇦 Ukrainian",
                "Bulgarian": "🇧🇬 Bulgarian",
                "Croatian": "🇭🇷 Croatian",
                "Serbian": "🇷🇸 Serbian",
                "Slovenian": "🇸🇮 Slovenian",
                "Slovak": "🇸🇰 Slovak",
                "Lithuanian": "🇱🇹 Lithuanian",
                "Latvian": "🇱🇻 Latvian",
                "Estonian": "🇪🇪 Estonian"
            }
            
            selected_language = st.selectbox(
                "Choose the language for slide explanations:",
                options=list(language_options.keys()),
                format_func=lambda x: language_options[x],
                index=0 if st.session_state.selected_language == "Spanish" else list(language_options.keys()).index(st.session_state.selected_language) if st.session_state.selected_language in language_options else 0,
                help="Select the language in which you want the AI to generate explanations"
            )
            
            # Update session state
            st.session_state.selected_language = selected_language
            
            st.markdown("")  # Add some space
            
            if st.button("🚀 Analyze All Slides", type="primary"):

                # Create progress tracking
                progress_bar = st.progress(0)
                status_text = st.empty()

                explanations = []

                for i, slide_bytes in enumerate(st.session_state.slides):
                    slide_num = i + 1
                    status_text.text(f"Analyzing slide {slide_num} of {len(st.session_state.slides)}...")

                    # Analyze slide
                    explanation = explain_slide(slide_bytes, openai_client, slide_num, custom_prompt, st.session_state.selected_language)
                    explanations.append(explanation)

                    # Update progress
                    progress_bar.progress((i + 1) / len(st.session_state.slides))

                status_text.text("✅ Analysis complete!")

                # Store results in session state
                st.session_state.explanations = explanations
                st.session_state.edited_explanations = [exp.copy() for exp in explanations]  # Initialize edited version
        
        # Display results if available
        if st.session_state.explanations is not None:
            st.markdown("**📋 Slide Analysis Results**")  # Changed from subheader to markdown for less space

            # Undo/Redo buttons with keyboard shortcuts
            col1, col2, col3 = st.columns([1, 1, 2])
            with col1:
                if st.button("↶ Undo", disabled=len(st.session_state.undo_stack) == 0):
                    if st.session_state.undo_stack:
                        # Save current state to redo stack
                        if st.session_state.slides and st.session_state.edited_explanations:
                            st.session_state.redo_stack.append({
                                'slides': st.session_state.slides.copy(),
                                'edited_explanations': [exp.copy() for exp in st.session_state.edited_explanations]
                            })
                        # Restore from undo stack
                        prev_state = st.session_state.undo_stack.pop()
                        st.session_state.slides = prev_state['slides']
                        st.session_state.edited_explanations = prev_state['edited_explanations']
                        st.rerun()

            with col2:
                if st.button("↷ Redo", disabled=len(st.session_state.redo_stack) == 0):
                    if st.session_state.redo_stack:
                        # Save current state to undo stack
                        if st.session_state.slides and st.session_state.edited_explanations:
                            st.session_state.undo_stack.append({
                                'slides': st.session_state.slides.copy(),
                                'edited_explanations': [exp.copy() for exp in st.session_state.edited_explanations]
                            })
                        # Restore from redo stack
                        next_state = st.session_state.redo_stack.pop()
                        st.session_state.slides = next_state['slides']
                        st.session_state.edited_explanations = next_state['edited_explanations']
                        st.rerun()

            with col3:
                st.caption("💡 Undo/Redo via buttons or Ctrl+Z / Ctrl+Shift+Z")

            # Add keyboard event handling for undo/redo
            st.markdown("""
            <script>
            document.addEventListener('keydown', function(event) {
                // Check for Ctrl+Z (undo)
                if (event.ctrlKey && event.key === 'z' && !event.shiftKey) {
                    event.preventDefault();
                    // Find and click the undo button
                    const undoBtn = document.querySelector('button[data-testid*="undo"]');
                    if (undoBtn && !undoBtn.disabled) {
                        undoBtn.click();
                    }
                }
                // Check for Ctrl+Shift+Z (redo)
                if (event.ctrlKey && event.shiftKey && event.key === 'Z') {
                    event.preventDefault();
                    // Find and click the redo button
                    const redoBtn = document.querySelector('button[data-testid*="redo"]');
                    if (redoBtn && !redoBtn.disabled) {
                        redoBtn.click();
                    }
                }
            });
            </script>
            """, unsafe_allow_html=True)

            # Use edited explanations if available, otherwise use original
            current_explanations = st.session_state.edited_explanations or st.session_state.explanations

            if not current_explanations or not st.session_state.slides:
                st.error("No explanations or slides available")
                st.stop()

            for i, (slide_bytes, explanation) in enumerate(zip(st.session_state.slides, current_explanations)):
                slide_num = i + 1

                with st.expander(f"📊 Slide {slide_num} Analysis", expanded=True):

                    # Slide controls
                    slide_col1, slide_col2, slide_col3 = st.columns([2, 1, 1])

                    with slide_col1:
                        # Click to enlarge slide
                        if st.button(f"🔍 Enlarge Slide {slide_num}", key=f"enlarge_{i}"):
                            st.session_state.current_slide_view = i
                            st.rerun()

                    with slide_col2:
                        # Edit button
                        if st.button(f"✏️ Edit Text", key=f"edit_{i}"):
                            st.session_state[f"edit_mode_{i}"] = not st.session_state.get(f"edit_mode_{i}", False)
                            st.rerun()

                    with slide_col3:
                        # Delete slide button
                        if st.button(f"🗑️ Delete", key=f"delete_{i}"):
                            # Save current state to undo stack
                            if st.session_state.slides and current_explanations:
                                st.session_state.undo_stack.append({
                                    'slides': st.session_state.slides.copy(),
                                    'edited_explanations': [exp.copy() for exp in current_explanations]
                                })
                            # Remove slide and explanation
                            st.session_state.slides.pop(i)
                            current_explanations.pop(i)
                            st.session_state.edited_explanations = current_explanations
                            # Clear redo stack
                            st.session_state.redo_stack = []
                            st.rerun()

                    # Display slide image (larger, full width) with darker background
                    st.markdown(f"""
                    <div style="background: rgba(0, 0, 0, 0.8); padding: 15px; border-radius: 10px; margin: 10px 0;">
                        <h3 style="color: #ffffff; margin-bottom: 10px;">🖼️ Slide {slide_num}</h3>
                    </div>
                    """, unsafe_allow_html=True)
                    st.image(slide_bytes, caption=f"Slide {slide_num}", width='stretch')

                    # Display explanation below the image with professional styling
                    st.markdown(f"""
                    <div style="background: linear-gradient(135deg, rgba(15,15,35,0.98), rgba(25,25,50,0.95), rgba(40,40,70,0.92)); padding: 25px; border-radius: 15px; box-shadow: 0 15px 40px rgba(0,0,0,0.4), inset 0 1px 0 rgba(255,255,255,0.1); backdrop-filter: blur(15px); border: 1px solid rgba(255,255,255,0.15); margin: 15px 0; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;">
                        <h3 style="color: #ffffff; margin-bottom: 20px; font-size: 1.8em; text-shadow: 2px 2px 4px rgba(0,0,0,0.5); font-weight: 700; border-bottom: 2px solid rgba(255,255,255,0.3); padding-bottom: 10px;">🔍 AI Analysis</h3>
                    """, unsafe_allow_html=True)

                    if explanation["success"]:
                        exp_data = explanation["explanation"]

                        # Check if in edit mode
                        edit_mode = st.session_state.get(f"edit_mode_{i}", False)

                        if edit_mode:
                            # Store initial values when entering edit mode (only once)
                            if f"initial_values_{i}" not in st.session_state:
                                st.session_state[f"initial_values_{i}"] = {
                                    'titulo': exp_data.get('titulo', ''),
                                    'explicacion_didactica': exp_data.get('explicacion_didactica', []),
                                    'puntos_clave': exp_data.get('puntos_clave', []),
                                    'conexiones': exp_data.get('conexiones', ''),
                                    'resumen_corto': exp_data.get('resumen_corto', '')
                                }
                                # Normalize explicacion_didactica
                                if isinstance(st.session_state[f"initial_values_{i}"]['explicacion_didactica'], str):
                                    st.session_state[f"initial_values_{i}"]['explicacion_didactica'] = [st.session_state[f"initial_values_{i}"]['explicacion_didactica']]

                                # Initialize form values from initial data
                                initial_vals = st.session_state[f"initial_values_{i}"]
                                st.session_state[f"title_{i}"] = initial_vals['titulo']
                                st.session_state[f"conexiones_{i}"] = initial_vals['conexiones']
                                st.session_state[f"resumen_corto_{i}"] = initial_vals['resumen_corto']

                                # Initialize explicacion inputs
                                for k, item in enumerate(initial_vals['explicacion_didactica']):
                                    item_id = f"id_{k}"
                                    st.session_state[f"explicacion_{i}_{item_id}"] = item

                                # Initialize puntos clave inputs
                                for k, item in enumerate(initial_vals['puntos_clave']):
                                    item_id = f"id_{k}"
                                    st.session_state[f"punto_{i}_{item_id}"] = item

                            # Editable fields with modern styling
                            st.markdown("""
                            <div style="background: linear-gradient(135deg, rgba(255,193,7,0.1), rgba(255,193,7,0.05)); padding: 20px; border-radius: 10px; border-left: 4px solid #FFC107; margin: 10px 0;">
                                <h4 style="color: #ffffff; margin-bottom: 15px; font-weight: 600;">✏️ Edit Mode</h4>
                            </div>
                            """, unsafe_allow_html=True)

                            # Back button to exit edit mode
                            if st.button("⬅️ Back to View", key=f"back_{i}"):
                                # Get initial values stored when entering edit mode
                                initial_values = st.session_state.get(f"initial_values_{i}", {})

                                # Check if there are unsaved changes by comparing current form state with initial values
                                changes_made = False

                                # Check title
                                initial_title = initial_values.get('titulo', '')
                                form_title = st.session_state.get(f"title_{i}", initial_title)
                                if form_title != initial_title:
                                    changes_made = True

                                # Check conexiones
                                initial_conexiones = initial_values.get('conexiones', '')
                                form_conexiones = st.session_state.get(f"conexiones_{i}", initial_conexiones)
                                if form_conexiones != initial_conexiones:
                                    changes_made = True

                                # Check resumen corto
                                initial_resumen = initial_values.get('resumen_corto', '')
                                form_resumen = st.session_state.get(f"resumen_corto_{i}", initial_resumen)
                                if form_resumen != initial_resumen:
                                    changes_made = True

                                # Check explicacion didactica - compare all possible form fields using IDs
                                initial_explicacion = initial_values.get('explicacion_didactica', [])

                                # Get all explicacion form values using IDs (check up to reasonable number)
                                form_explicacion_values = []
                                explicacion_ids = st.session_state.get(f"explicacion_ids_{i}", [])
                                for item_id in explicacion_ids:
                                    key = f"explicacion_{i}_{item_id}"
                                    if key in st.session_state:
                                        form_explicacion_values.append(st.session_state[key])

                                # Check for changes: length differences (additions/removals) or content differences
                                if len(form_explicacion_values) != len(initial_explicacion):
                                    changes_made = True
                                else:
                                    # Filter out empty strings from both for content comparison
                                    initial_explicacion_filtered = [x for x in initial_explicacion if x.strip()]
                                    form_explicacion_filtered = [x for x in form_explicacion_values if x.strip()]
                                    if form_explicacion_filtered != initial_explicacion_filtered:
                                        changes_made = True

                                # Check puntos clave - compare all possible form fields using IDs
                                initial_puntos = initial_values.get('puntos_clave', [])

                                # Get all puntos form values using IDs (check up to reasonable number)
                                form_puntos_values = []
                                puntos_ids = st.session_state.get(f"puntos_ids_{i}", [])
                                for item_id in puntos_ids:
                                    key = f"punto_{i}_{item_id}"
                                    if key in st.session_state:
                                        form_puntos_values.append(st.session_state[key])

                                # Check for changes: length differences (additions/removals) or content differences
                                if len(form_puntos_values) != len(initial_puntos):
                                    changes_made = True
                                else:
                                    # Filter out empty strings from both for content comparison
                                    initial_puntos_filtered = [x for x in initial_puntos if x.strip()]
                                    form_puntos_filtered = [x for x in form_puntos_values if x.strip()]
                                    if form_puntos_filtered != initial_puntos_filtered:
                                        changes_made = True

                                if changes_made:
                                    # Set flag to show confirmation dialog
                                    st.session_state[f"show_confirm_exit_{i}"] = True
                                else:
                                    # No changes, exit directly and clear initial values
                                    st.session_state[f"edit_mode_{i}"] = False
                                    if f"initial_values_{i}" in st.session_state:
                                        del st.session_state[f"initial_values_{i}"]
                                    if f"puntos_ids_{i}" in st.session_state:
                                        del st.session_state[f"puntos_ids_{i}"]
                                    st.rerun()

                            # Show confirmation dialog if flag is set
                            if st.session_state.get(f"show_confirm_exit_{i}", False):
                                st.warning("⚠️ You have unsaved changes. Are you sure you want to exit without saving?")
                                col1, col2 = st.columns(2)
                                with col1:
                                    if st.button("Yes, Exit Without Saving", key=f"confirm_exit_{i}"):
                                        # Restore initial values to exp_data to discard all changes
                                        initial_values = st.session_state.get(f"initial_values_{i}", {})
                                        exp_data.update(initial_values)

                                        # Clear edit-related session state
                                        keys_to_clear = [f"initial_values_{i}", f"show_confirm_exit_{i}"]
                                        if f"puntos_ids_{i}" in st.session_state:
                                            keys_to_clear.append(f"puntos_ids_{i}")
                                        if f"explicacion_ids_{i}" in st.session_state:
                                            keys_to_clear.append(f"explicacion_ids_{i}")

                                        for key in keys_to_clear:
                                            if key in st.session_state:
                                                del st.session_state[key]

                                        st.session_state[f"edit_mode_{i}"] = False
                                        st.rerun()
                                with col2:
                                    if st.button("No, Stay in Edit Mode", key=f"cancel_exit_{i}"):
                                        st.session_state[f"show_confirm_exit_{i}"] = False
                                        st.rerun()

                            # Title
                            new_title = st.text_input(
                                "📌 Título:",
                                value=exp_data.get('titulo', ''),
                                key=f"title_{i}"
                            )

                            # Explicacion didactica - individual text inputs with drag handles (up/down arrows)
                            st.markdown("**🧠 Explicación didáctica:**")
                            explicacion_list = exp_data.get('explicacion_didactica', [])
                            if isinstance(explicacion_list, str):
                                explicacion_list = [explicacion_list]
                            elif not isinstance(explicacion_list, list):
                                explicacion_list = []

                            # Assign unique IDs to items if not present
                            if f"explicacion_ids_{i}" not in st.session_state:
                                st.session_state[f"explicacion_ids_{i}"] = [f"id_{k}" for k in range(len(explicacion_list))]
                            explicacion_ids = st.session_state[f"explicacion_ids_{i}"]

                            # Ensure IDs match list length
                            while len(explicacion_ids) < len(explicacion_list):
                                explicacion_ids.append(f"id_{len(explicacion_ids)}")
                            while len(explicacion_ids) > len(explicacion_list):
                                explicacion_ids.pop()

                            new_explicacion_didactica = []
                            for j, (item, item_id) in enumerate(zip(explicacion_list, explicacion_ids)):
                                col1, col2, col3 = st.columns([3, 1, 1])
                                with col1:
                                    new_item = st.text_input(
                                        f"Punto {j+1}:",
                                        value=item,
                                        key=f"explicacion_{i}_{item_id}",
                                        label_visibility="collapsed"
                                    )
                                with col2:
                                    # Drag handles (up/down arrows)
                                    arrow_col1, arrow_col2 = st.columns(2)
                                    with arrow_col1:
                                        if j > 0 and st.button("⬆️", key=f"up_explicacion_{i}_{item_id}"):
                                            # Swap items and IDs
                                            explicacion_list[j], explicacion_list[j-1] = explicacion_list[j-1], explicacion_list[j]
                                            explicacion_ids[j], explicacion_ids[j-1] = explicacion_ids[j-1], explicacion_ids[j]
                                            st.rerun()
                                    with arrow_col2:
                                        if j < len(explicacion_list) - 1 and st.button("⬇️", key=f"down_explicacion_{i}_{item_id}"):
                                            # Swap items and IDs
                                            explicacion_list[j], explicacion_list[j+1] = explicacion_list[j+1], explicacion_list[j]
                                            explicacion_ids[j], explicacion_ids[j+1] = explicacion_ids[j+1], explicacion_ids[j]
                                            st.rerun()
                                with col3:
                                    if st.button("🗑️", key=f"remove_explicacion_{i}_{item_id}"):
                                        # Remove item and ID by index j
                                        explicacion_list.pop(j)
                                        explicacion_ids.pop(j)
                                        st.rerun()
                                new_explicacion_didactica.append(new_item)

                            # Add new point button
                            if st.button("➕ Add Point", key=f"add_point_{i}"):
                                explicacion_list.append("")
                                explicacion_ids.append(f"id_{len(explicacion_ids)}")
                                st.rerun()

                            # Puntos clave - individual text inputs for each point with drag handles (up/down arrows)
                            st.markdown("**🎯 Puntos clave:**")
                            puntos_clave_list = exp_data.get('puntos_clave', [])
                            if not isinstance(puntos_clave_list, list):
                                puntos_clave_list = []

                            # Assign unique IDs to items if not present
                            if f"puntos_ids_{i}" not in st.session_state:
                                st.session_state[f"puntos_ids_{i}"] = [f"id_{k}" for k in range(len(puntos_clave_list))]
                            puntos_ids = st.session_state[f"puntos_ids_{i}"]

                            # Ensure IDs match list length
                            while len(puntos_ids) < len(puntos_clave_list):
                                puntos_ids.append(f"id_{len(puntos_ids)}")
                            while len(puntos_ids) > len(puntos_clave_list):
                                puntos_ids.pop()

                            new_puntos_clave = []
                            for j, (item, item_id) in enumerate(zip(puntos_clave_list, puntos_ids)):
                                col1, col2, col3 = st.columns([3, 1, 1])
                                with col1:
                                    new_item = st.text_input(
                                        f"Item {j+1}:",
                                        value=item,
                                        key=f"punto_{i}_{item_id}",
                                        label_visibility="collapsed"
                                    )
                                with col2:
                                    # Drag handles (up/down arrows)
                                    arrow_col1, arrow_col2 = st.columns(2)
                                    with arrow_col1:
                                        if j > 0 and st.button("⬆️", key=f"up_punto_{i}_{item_id}"):
                                            # Swap items and IDs
                                            puntos_clave_list[j], puntos_clave_list[j-1] = puntos_clave_list[j-1], puntos_clave_list[j]
                                            puntos_ids[j], puntos_ids[j-1] = puntos_ids[j-1], puntos_ids[j]
                                            st.rerun()
                                    with arrow_col2:
                                        if j < len(puntos_clave_list) - 1 and st.button("⬇️", key=f"down_punto_{i}_{item_id}"):
                                            # Swap items and IDs
                                            puntos_clave_list[j], puntos_clave_list[j+1] = puntos_clave_list[j+1], puntos_clave_list[j]
                                            puntos_ids[j], puntos_ids[j+1] = puntos_ids[j+1], puntos_ids[j]
                                            st.rerun()
                                with col3:
                                    if st.button("🗑️", key=f"remove_punto_{i}_{item_id}"):
                                        # Remove item and ID by index j
                                        puntos_clave_list.pop(j)
                                        puntos_ids.pop(j)
                                        st.rerun()
                                new_puntos_clave.append(new_item)

                            # Add new point button for puntos clave
                            if st.button("➕ Add Point", key=f"add_punto_{i}"):
                                puntos_clave_list.append("")
                                puntos_ids.append(f"id_{len(puntos_ids)}")
                                st.rerun()

                            # Update the puntos_clave in exp_data with the edited values
                            exp_data['puntos_clave'] = new_puntos_clave

                            # Conexiones
                            new_conexiones = st.text_area(
                                "🔗 Conexiones:",
                                value=exp_data.get('conexiones', ''),
                                key=f"conexiones_{i}",
                                height=80
                            )

                            # Resumen corto
                            new_resumen_corto = st.text_area(
                                "📝 Resumen corto:",
                                value=exp_data.get('resumen_corto', ''),
                                key=f"resumen_corto_{i}",
                                height=60
                            )

                            # Save button
                            if st.button("💾 Save Changes", key=f"save_{i}"):
                                # Save current state to undo stack
                                if st.session_state.slides and current_explanations:
                                    st.session_state.undo_stack.append({
                                        'slides': st.session_state.slides.copy(),
                                        'edited_explanations': [exp.copy() for exp in current_explanations]
                                    })

                                # Process explicacion_didactica - filter out empty points
                                processed_explicacion = [item for item in new_explicacion_didactica if item.strip()]
                                if len(processed_explicacion) == 1:
                                    processed_explicacion = processed_explicacion[0]
                                elif len(processed_explicacion) == 0:
                                    processed_explicacion = ""

                                # Update the explanation with new schema
                                updated_exp = explanation.copy()
                                updated_exp["explanation"] = {
                                    'titulo': new_title,
                                    'explicacion_didactica': processed_explicacion,
                                    'puntos_clave': new_puntos_clave,
                                    'conexiones': new_conexiones,
                                    'resumen_corto': new_resumen_corto
                                }

                                # Update edited explanations
                                if st.session_state.edited_explanations is None:
                                    st.session_state.edited_explanations = [exp.copy() for exp in st.session_state.explanations]
                                st.session_state.edited_explanations[i] = updated_exp

                                # Clear redo stack
                                st.session_state.redo_stack = []

                                # Exit edit mode
                                st.session_state[f"edit_mode_{i}"] = False
                                st.success("✅ Changes saved!")
                                st.rerun()

                        else:
                            # Display mode with professional Word-like styling (content inside AI Analysis container)

                            # Title section
                            titulo_ui = exp_data.get('titulo', 'N/A')
                            if not use_custom_prompt:
                                st.markdown(f"<h4 style='color: #ffffff; margin-bottom: 15px; font-weight: 600; border-bottom: 1px solid rgba(255,255,255,0.3); padding-bottom: 8px;'><span style='color: #4CAF50;'>📌</span> Título: <strong>{titulo_ui}</strong></h4>", unsafe_allow_html=True)
                            else:
                                st.markdown(f"<h4 style='color: #ffffff; margin-bottom: 15px; font-weight: 600;'><strong>📌 Título:</strong> {titulo_ui}</h4>", unsafe_allow_html=True)

                            # Explicación didáctica section
                            explicacion_ui = exp_data.get('explicacion_didactica') or exp_data.get('resumen') or 'N/A'
                            st.markdown("<h4 style='color: #ffffff; margin-bottom: 12px; font-weight: 600;'><span style='color: #2196F3;'>🧠</span> Explicación didáctica:</h4>", unsafe_allow_html=True)
                            if isinstance(explicacion_ui, list):
                                for punto in explicacion_ui:
                                    st.markdown(f"<p style='color: #ffffff; line-height: 1.6; margin-bottom: 8px; margin-left: 20px; text-indent: -15px;'><span style='color: #cccccc; margin-right: 10px;'>•</span>{punto}</p>", unsafe_allow_html=True)
                            else:
                                st.markdown(f"<p style='color: #ffffff; line-height: 1.6; margin-bottom: 15px;'>{explicacion_ui}</p>", unsafe_allow_html=True)

                            # Puntos clave section
                            puntos_ui = exp_data.get('puntos_clave') or exp_data.get('contenido_clave') or []
                            if puntos_ui:
                                st.markdown("<h4 style='color: #ffffff; margin-bottom: 12px; margin-top: 20px; font-weight: 600;'><span style='color: #FF9800;'>🎯</span> Puntos clave:</h4>", unsafe_allow_html=True)
                                for item in puntos_ui:
                                    st.markdown(f"<p style='color: #ffffff; line-height: 1.6; margin-bottom: 6px; margin-left: 20px; text-indent: -15px;'><span style='color: #cccccc; margin-right: 10px;'>-</span>{item}</p>", unsafe_allow_html=True)

                            # Conexiones section
                            conex_ui = exp_data.get('conexiones') or exp_data.get('contexto') or ''
                            if conex_ui:
                                st.markdown("<h4 style='color: #ffffff; margin-bottom: 12px; margin-top: 20px; font-weight: 600;'><span style='color: #9C27B0;'>🔗</span> Conexiones:</h4>", unsafe_allow_html=True)
                                st.markdown(f"<p style='color: #ffffff; line-height: 1.6; margin-bottom: 15px;'>{conex_ui}</p>", unsafe_allow_html=True)

                            # Resumen corto section
                            resumen_corto_ui = exp_data.get('resumen_corto') or exp_data.get('resumen') or ''
                            if resumen_corto_ui:
                                st.markdown("<h4 style='color: #ffffff; margin-bottom: 12px; margin-top: 20px; font-weight: 600;'><span style='color: #607D8B;'>📝</span> Resumen corto:</h4>", unsafe_allow_html=True)
                                st.markdown(f"<p style='color: #ffffff; line-height: 1.6; margin-bottom: 15px; font-style: italic;'>{resumen_corto_ui}</p>", unsafe_allow_html=True)

                            # Anki cards section
                            anki_cards_ui = exp_data.get('anki_cards') or []
                            if anki_cards_ui:
                                st.markdown("<h4 style='color: #ffffff; margin-bottom: 12px; margin-top: 20px; font-weight: 600;'><span style='color: #FFA726;'>🃏</span> Tarjetas Anki:</h4>", unsafe_allow_html=True)
                                for idx, card in enumerate(anki_cards_ui, 1):
                                    if isinstance(card, dict) and 'pregunta' in card and 'respuesta' in card:
                                        st.markdown(f"""
                                        <div style='background: linear-gradient(135deg, rgba(255,167,38,0.1), rgba(255,167,38,0.05)); 
                                                    border-left: 3px solid #FFA726; padding: 15px; margin: 10px 0; border-radius: 8px;'>
                                            <p style='color: #FFA726; font-weight: 600; margin-bottom: 8px;'>📋 Pregunta {idx}:</p>
                                            <p style='color: #ffffff; margin-bottom: 12px; font-style: italic;'>{card['pregunta']}</p>
                                            <p style='color: #FFA726; font-weight: 600; margin-bottom: 8px;'>💡 Respuesta:</p>
                                            <p style='color: #ffffff; margin-bottom: 0;'>{card['respuesta']}</p>
                                        </div>
                                        """, unsafe_allow_html=True)

                            # Insights section
                            insights_ui = exp_data.get('insights') or []
                            if insights_ui:
                                st.markdown("<h4 style='color: #ffffff; margin-bottom: 12px; margin-top: 20px; font-weight: 600;'><span style='color: #00BCD4;'>💡</span> Insights:</h4>", unsafe_allow_html=True)
                                for item in insights_ui:
                                    st.markdown(f"<p style='color: #ffffff; line-height: 1.6; margin-bottom: 6px; margin-left: 20px; text-indent: -15px;'><span style='color: #cccccc; margin-right: 10px;'>-</span>{item}</p>", unsafe_allow_html=True)

                    else:
                        st.error(f"❌ {explanation.get('error', 'Unknown error')}")
                        st.markdown("</div>", unsafe_allow_html=True)  # Close the AI Analysis div

            # Add separator between slides for better separation
            st.markdown("---")
            st.markdown("")
            st.markdown("")  # Extra space between slides

            # Modal for enlarged slide view
            if st.session_state.current_slide_view is not None:
                i = st.session_state.current_slide_view
                slide_bytes = st.session_state.slides[i]
                slide_num = i + 1

                # Make the overlay clickable to close (no visible close button)
                st.markdown("""
                <style>
                .enlarged-slide {
                    cursor: pointer;
                }
                .enlarged-slide img {
                    pointer-events: none; /* Allow clicks on the overlay to pass through to close */
                }
                </style>
                """, unsafe_allow_html=True)

                # Full screen enlarged image with custom styling
                st.markdown("""
                <style>
                .enlarged-slide {
                    position: fixed;
                    top: 0;
                    left: 0;
                    width: 100vw;
                    height: 100vh;
                    background: rgba(0,0,0,0.9);
                    z-index: 9999;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    padding: 20px;
                    box-sizing: border-box;
                }
                .enlarged-slide img {
                    max-width: 95vw;
                    max-height: 90vh;
                    object-fit: contain;
                    border-radius: 10px;
                    box-shadow: 0 0 50px rgba(0,0,0,0.5);
                }
                </style>
                """, unsafe_allow_html=True)

                # Create a temporary image file for full-screen display
                import tempfile
                import os
                from PIL import Image as PILImage
                import io

                # Convert bytes to PIL Image
                image = PILImage.open(io.BytesIO(slide_bytes))

                # Save to temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                    image.save(tmp_file.name, 'PNG')
                    temp_image_path = tmp_file.name

                # Close button at top right
                col1, col2 = st.columns([10, 1])
                with col2:
                    if st.button("❌ Close", key="close_enlarged"):
                        st.session_state.current_slide_view = None
                        st.rerun()

                # Full width enlarged image
                st.image(slide_bytes, caption=f"Slide {slide_num} (Enlarged)", use_container_width=True)

                # Clean up temp file
                try:
                    os.unlink(temp_image_path)
                except:
                    pass
            
            # Export options
            st.markdown("**📤 Export Results**")  # Changed from subheader to markdown for less space

            col1, col2, col3 = st.columns(3)

            with col1:
                # Word Document Export
                if st.button("📄 Generate Word Report", key="generate_word"):
                    with st.spinner("🔄 Generating Word document... This may take a moment depending on the number of slides."):
                        try:
                            # Use edited explanations if available, otherwise use original
                            explanations_to_use = st.session_state.edited_explanations or st.session_state.explanations
                            st.session_state.word_report = generate_word_report(
                                st.session_state.slides,
                                explanations_to_use,
                                st.session_state.uploaded_file_name
                            )
                            st.success("✅ Word document generated successfully!")

                        except Exception as e:
                            st.error(f"❌ Error generating Word document: {str(e)}")

                # Show Word preview and download button if report is generated
                if st.session_state.word_report is not None:
                    # Enhanced Preview section
                    with st.expander("📋 Word Report Preview", expanded=False):
                        try:
                            # Extract structured content from the generated Word document for better preview
                            import io

                            # Load the Word document from bytes
                            from docx import Document as DocxDocument
                            doc = DocxDocument(io.BytesIO(st.session_state.word_report))

                            # Create a structured preview with formatting
                            preview_sections = []

                            for para in doc.paragraphs:
                                text = para.text.strip()
                                if text:
                                    # Identify section headers and format accordingly
                                    if any(header in text.upper() for header in ["TÍTULO", "EXPLICACIÓN", "PUNTOS", "CONEXIONES", "RESUMEN"]):
                                        preview_sections.append(f"**{text}**")
                                    elif text.startswith("•") or text.startswith("-"):
                                        preview_sections.append(f"  {text}")
                                    else:
                                        preview_sections.append(text)

                            # Limit preview to reasonable length but show complete sections
                            preview_content = "\n\n".join(preview_sections)
                            max_length = 1500

                            if len(preview_content) > max_length:
                                # Try to cut at a section boundary
                                truncated = preview_content[:max_length]
                                last_section_end = max(
                                    truncated.rfind("\n\n**"),
                                    truncated.rfind("\n\n")
                                )
                                if last_section_end > max_length * 0.7:  # If we can cut at a reasonable boundary
                                    preview_content = preview_content[:last_section_end]
                                else:
                                    preview_content = preview_content[:max_length]

                                preview_content += "\n\n[... Preview truncated - full document available for download ...]"

                            # Display with better formatting
                            st.markdown("### 📄 Document Preview")
                            st.markdown("---")

                            # Show document info
                            total_slides = len(st.session_state.slides)
                            st.info(f"📊 This report contains analysis for {total_slides} slides with detailed explanations, key points, and connections.")

                            # Display preview in a styled container
                            st.markdown("""
                            <div style="background: rgba(255, 255, 255, 0.95); padding: 20px; border-radius: 10px; border-left: 4px solid #4CAF50; margin: 10px 0;">
                            """, unsafe_allow_html=True)

                            # Split content and display with proper formatting
                            lines = preview_content.split('\n')
                            for line in lines:
                                if line.strip():
                                    if '**' in line and line.count('**') >= 2:
                                        # Section headers
                                        st.markdown(f"#### {line.strip('*')}")
                                    elif line.startswith('  •') or line.startswith('  -'):
                                        # Bullet points
                                        st.markdown(line.strip())
                                    else:
                                        # Regular text
                                        st.write(line.strip())

                            st.markdown("</div>", unsafe_allow_html=True)

                            # Show file size info
                            file_size_kb = len(st.session_state.word_report) / 1024
                            st.caption(f"📁 File size: {file_size_kb:.1f} KB")

                        except Exception as e:
                            st.error(f"❌ Error generating preview: {str(e)}")
                            st.info("💡 The document was generated successfully, but preview failed. You can still download the full report.")

                    # Download button
                    if st.session_state.uploaded_file_name:
                        st.download_button(
                            label="📥 Download Word Report",
                            data=st.session_state.word_report,
                            file_name=f"{st.session_state.uploaded_file_name.replace('.pdf', '')}_analysis_report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="word_download"
                        )
                    else:
                        st.download_button(
                            label="📥 Download Word Report",
                            data=st.session_state.word_report,
                            file_name="analysis_report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="word_download"
                        )

            with col2:
                # Anki Cards Export
                if st.button("🃏 Generate Anki Cards", key="generate_anki"):
                    with st.spinner("🔄 Generating Anki cards..."):
                        try:
                            # Use edited explanations if available, otherwise use original
                            explanations_to_use = st.session_state.edited_explanations or st.session_state.explanations
                            anki_content = generate_anki_export(
                                explanations_to_use,
                                st.session_state.uploaded_file_name
                            )
                            st.session_state.anki_cards_export = anki_content
                            st.success("✅ Anki cards generated successfully!")
                        except Exception as e:
                            st.error(f"❌ Error generating Anki cards: {str(e)}")

                # Show Anki download button if cards are generated
                if hasattr(st.session_state, 'anki_cards_export') and st.session_state.anki_cards_export:
                    if st.session_state.uploaded_file_name:
                        st.download_button(
                            label="📥 Download Anki Cards",
                            data=st.session_state.anki_cards_export,
                            file_name=f"{st.session_state.uploaded_file_name.replace('.pdf', '')}_anki_cards.apkg",
                            mime="application/octet-stream",
                            key="anki_download"
                        )
                    else:
                        st.download_button(
                            label="📥 Download Anki Cards",
                            data=st.session_state.anki_cards_export,
                            file_name="anki_cards.apkg",
                            mime="application/octet-stream",
                            key="anki_download"
                        )

            with col3:
                # Quiz Generation
                if st.button("🧠 Generate Quiz", key="generate_quiz"):
                    with st.spinner("🔄 Generating quiz..."):
                        try:
                            # Use edited explanations if available, otherwise use original
                            explanations_to_use = st.session_state.edited_explanations or st.session_state.explanations

                            # Collect all anki cards
                            all_anki_cards = []
                            for explanation in explanations_to_use:
                                if explanation.get("success") and explanation.get("explanation"):
                                    anki_cards = explanation["explanation"].get("anki_cards", [])
                                    all_anki_cards.extend(anki_cards)

                            if not all_anki_cards:
                                st.error("❌ No Anki cards available to generate quiz")
                            else:
                                quiz_questions = generate_quiz(all_anki_cards)
                                if not quiz_questions:
                                    st.error("❌ Not enough Anki cards to generate quiz (need at least 4)")
                                else:
                                    # Clear previous quiz answer states
                                    keys_to_clear = [key for key in st.session_state.keys() if isinstance(key, str) and (key.startswith('quiz_answer_selected_') or key.startswith('quiz_selected_option_'))]
                                    for key in keys_to_clear:
                                        del st.session_state[key]

                                    st.session_state.quiz_questions = quiz_questions
                                    st.session_state.quiz_current_index = 0
                                    st.session_state.quiz_score = 0
                                    st.session_state.quiz_show_feedback = False
                                    st.success(f"✅ Quiz generated with {len(quiz_questions)} questions!")

                        except Exception as e:
                            st.error(f"❌ Error generating quiz: {str(e)}")

            # Quiz Section (below all exports)
            if 'quiz_questions' in st.session_state and st.session_state.quiz_questions:
                st.markdown("---")
                st.markdown("## 🧠 Interactive Quiz")

                quiz_questions = st.session_state.quiz_questions
                current_index = st.session_state.quiz_current_index

                if current_index < len(quiz_questions):
                    question = quiz_questions[current_index]

                    st.markdown(f"**Question {current_index + 1} of {len(quiz_questions)}**")
                    st.markdown(f"### {question['question']}")

                    # Options as buttons in 2 columns
                    cols = st.columns(2)
                    selected_option = None

                    # Check if answer was already selected for this question
                    answer_selected = st.session_state.get(f"quiz_answer_selected_{current_index}", False)
                    selected_option = st.session_state.get(f"quiz_selected_option_{current_index}", None)
                    is_correct = selected_option == question['correct_answer'] if selected_option else None

                    for i, option in enumerate(question['options']):
                        col_idx = i % 2
                        with cols[col_idx]:
                            # Disable buttons if answer was already selected
                            button_disabled = answer_selected

                            # Create button label with visual feedback
                            button_label = f"**{chr(65+i)}) {option}**"

                            if answer_selected:
                                if option == question['correct_answer']:
                                    button_label = f"**{chr(65+i)}) {option}** ✅"
                                elif option == selected_option and not is_correct:
                                    button_label = f"**{chr(65+i)}) {option}** ❌"

                            if st.button(button_label, key=f"quiz_option_{i}_{current_index}", disabled=button_disabled):
                                if not answer_selected:  # Only process if not already answered
                                    selected_option = option
                                    st.session_state[f"quiz_selected_option_{current_index}"] = selected_option
                                    st.session_state[f"quiz_answer_selected_{current_index}"] = True
                                    st.rerun()  # Immediate re-run to disable buttons instantly

                    if selected_option:
                        if is_correct:
                            st.session_state.quiz_score += 1
                            st.success(f"Correct! The answer is: {question['correct_answer']}", icon="✅")
                        else:
                            st.error(f"Incorrect. The correct answer is: {question['correct_answer']}", icon="❌")

                        # Wait 3 seconds
                        import time
                        time.sleep(3)

                        # Move to next question
                        st.session_state.quiz_current_index += 1
                        st.rerun()

                else:
                    # Quiz finished
                    score = st.session_state.quiz_score
                    total = len(quiz_questions)
                    percentage = (score / total) * 100

                    st.markdown("## 🎉 Quiz Completed!")
                    st.markdown(f"**Score: {score}/{total} ({percentage:.1f}%)**")

                    if percentage >= 80:
                        st.success("Excellent work! 🎊")
                    elif percentage >= 60:
                        st.info("Good work, keep practicing 📚")
                    else:
                        st.warning("You need more practice. You can do it! 💪")

                    if st.button("🔄 Take Quiz Again"):
                        # Reset quiz and clear all quiz-related session state
                        keys_to_clear = [key for key in st.session_state.keys() if isinstance(key, str) and (
                            key.startswith('quiz_answer_selected_') or
                            key.startswith('quiz_selected_option_')
                        )]
                        for key in keys_to_clear:
                            del st.session_state[key]

                        # Re-initialize quiz state but keep the questions
                        st.session_state.quiz_current_index = 0
                        st.session_state.quiz_score = 0
                        st.session_state.quiz_show_feedback = False
                        st.rerun()

if __name__ == "__main__":
    main()