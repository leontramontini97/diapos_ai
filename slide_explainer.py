"""PDF Slide Explainer using Flora Facturacion Infrastructure
Streamlit app that processes PDF slides and generates explanations for each slide
"""

import streamlit as st
import os
import base64
import json
from typing import List, Dict, Any
from openai import OpenAI
import fitz  # PyMuPDF for PDF processing
from PIL import Image
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import tempfile
import re


# === Prompt template por defecto ===
PROMPT_TEMPLATE = """\
Hazme una explicación **completa, clara y dinámica** sobre este texto.
Debe permitir al lector **entender todo el contenido técnico de manera fácil y ordenada**,
sin extenderse demasiado ni omitir ningún detalle importante.
Incluye ejemplos o analogías cuando ayuden a comprender mejor.
Al final, agrega un **resumen corto** con lo más importante de toda la explicación.

OBJETIVO GENERAL
- Que sea **profunda pero comprensible**, con rigor técnico y tono didáctico.
- Que ayude a **aprender de forma rápida**.
- Que combine explicación fluida con **puntos clave**.

INSTRUCCIONES
1) Explica el tema principal y por qué es relevante.
2) Explica cada punto técnico con claridad (mantén términos en inglés o alemán si aplica) y usa puntos clave para una explicación comprensible y comleta.
3) Resume conceptos principales en puntos clave.
4) Conecta con temas relacionados, pero haciéndolo específico y en relación con las demás diapositivas, no tan general. Aporta información realmente útil y que ayude a comprender mejor el tema, no datos innecesarios.
5) Cierra con un **resumen corto** (2–3 frases con el takeaway).

FORMATO DE SALIDA
Devuelve **únicamente** un **objeto JSON válido** (sin texto adicional, sin comentarios).
NO copies literalmente el ejemplo; rellénalo con el contenido del slide.

```json
{
  "titulo": "Tema o concepto central de la diapositiva",
  "explicacion_didactica": "Explicación completa y clara...",
  "puntos_clave": ["Idea 1", "Idea 2", "Idea 3"],
  "conexiones": "Relaciones con otros temas importantes",
  "resumen_corto": "Síntesis breve (2–3 frases)"
}
"""


def encode_image_base64(image_bytes: bytes) -> str:
    """Encode image bytes to base64 string"""
    return base64.b64encode(image_bytes).decode('utf-8')

def init_openai_client(api_key: str = None):
    """Initialize OpenAI client with API key"""
    if not api_key:
        api_key = os.getenv("OPENAI_API_KEY")
    
    if not api_key:
        return None
    
    try:
        return OpenAI(api_key=api_key)
    except Exception as e:
        st.error(f"❌ Error initializing OpenAI client: {str(e)}")
        return None

def extract_slides_from_pdf(pdf_file) -> List[bytes]:
    """
    Extract individual slides/pages from PDF as images
    
    Args:
        pdf_file: Uploaded PDF file from Streamlit
        
    Returns:
        List of image bytes for each slide
    """
    slides = []
    
    try:
        # Open PDF from uploaded file
        pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
        
        for page_num in range(len(pdf_document)):
            # Get page
            page = pdf_document.load_page(page_num)
            
            # Convert to image (300 DPI for good quality)
            mat = fitz.Matrix(300/72, 300/72)  # 300 DPI scaling
            pix = page.get_pixmap(matrix=mat)
            
            # Convert to PNG bytes
            img_data = pix.tobytes("png")
            slides.append(img_data)
            
        pdf_document.close()
        return slides
        
    except Exception as e:
        st.error(f"Error extracting slides from PDF: {str(e)}")
        return []

def explain_slide(slide_image_bytes: bytes, openai_client: OpenAI, slide_number: int, custom_prompt: str = None, language: str = "Spanish") -> Dict[str, Any]:
    """
    Generate explanation for a single slide using OpenAI Vision API
    
    Args:
        slide_image_bytes: Image bytes of the slide
        openai_client: OpenAI client instance
        slide_number: Number of the slide (for context)
        
    Returns:
        Dictionary with slide explanation and analysis
    """
    try:
        # Encode image to base64
        image_base64 = encode_image_base64(slide_image_bytes)
        image_url = f"data:image/png;base64,{image_base64}"
        
        # Use custom prompt if provided, otherwise use default
        # IMPORTANT: Avoid str.format here because PROMPT_TEMPLATE contains JSON braces
        # which would be interpreted as format fields. We only want to substitute {slide_number}.
        if custom_prompt:
            explanation_prompt = custom_prompt.replace("{slide_number}", str(slide_number))
        else:
            explanation_prompt = PROMPT_TEMPLATE.replace("{slide_number}", str(slide_number))
        
        # Add language instruction
        if language.lower() != "spanish":
            explanation_prompt += f"\n\nIMPORTANT: Provide the output in {language}."

        
        # Call Vision API
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            response_format={"type": "json_object"},  # <— NUEVO: fuerza JSON puro
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": explanation_prompt},
                        {"type": "image_url", "image_url": {"url": image_url, "detail": "high"}}
                    ]
                }
            ],
            max_tokens=2000,
            temperature=0  # <— recomendado para consistencia
        )
        
        # Parse response (normaliza a string y extrae JSON de forma robusta)
        raw = response.choices[0].message.content

        # Algunos SDK/dev builds pueden devolver None o listas de partes
        if raw is None:
            # intenta recuperar de un posible atributo alternativo
            raw = getattr(response.choices[0].message, "parsed", None)

        if isinstance(raw, list):
            content = "".join(
                (p.get("text", "") if isinstance(p, dict) else str(p)) for p in raw
            )
        elif raw is None:
            content = ""
        else:
            content = str(raw)

        def extract_json_safe(s: str):
            # Clean input first
            s = s.strip()
            
            # 1) JSON puro
            try:
                return json.loads(s)
            except Exception:
                pass
                
            # 2) Try to fix malformed JSON that starts with quotes
            if s.startswith('"') and not s.startswith('{"'):
                try:
                    # Try adding opening brace
                    fixed = "{" + s
                    return json.loads(fixed)
                except Exception:
                    try:
                        # Try removing leading quotes and finding JSON-like content
                        cleaned = s.lstrip('\n "')
                        if ':' in cleaned:
                            fixed = '{"' + cleaned
                            return json.loads(fixed)
                    except Exception:
                        pass
                    
            # 3) bloque ```json ... ```
            m = re.search(r"```json\s*(\{.*?\})\s*```", s, re.S)
            if m:
                try:
                    return json.loads(m.group(1))
                except Exception:
                    pass
                    
            # 4) bloque ``` ... ```
            m = re.search(r"```\s*(\{.*?\})\s*```", s, re.S)
            if m:
                try:
                    return json.loads(m.group(1))
                except Exception:
                    pass
                    
            # 5) primer objeto { ... }
            m = re.search(r"(\{.*\})", s, re.S)
            if m:
                try:
                    return json.loads(m.group(1))
                except Exception:
                    pass
                    
            # 6) último recurso: limpia ecos de {{ ... }} y usa como explicación
            cleaned = re.sub(r"\{\{[\s\S]*?\}\}", "", s).strip()
            return {
                "titulo": f"Slide {slide_number}",
                "explicacion_didactica": cleaned if cleaned else s,
                "puntos_clave": [],
                "conexiones": "",
                "resumen_corto": ""
            }

        explanation_data = extract_json_safe(content)

        
        # === Normalización de esquema al nuevo formato ===
        # Si ya viene en el esquema nuevo, lo usamos tal cual:
        if all(k in explanation_data for k in ["titulo", "explicacion_didactica", "puntos_clave", "conexiones", "resumen_corto"]):
            normalized = {
                "titulo": explanation_data.get("titulo", ""),
                "explicacion_didactica": explanation_data.get("explicacion_didactica", ""),
                "puntos_clave": explanation_data.get("puntos_clave", []) or [],
                "conexiones": explanation_data.get("conexiones", ""),
                "resumen_corto": explanation_data.get("resumen_corto", "")
            }
        else:
            # Fallback desde el esquema antiguo
            titulo_old = explanation_data.get("titulo", f"Slide {slide_number}")
            contenido_clave_old = explanation_data.get("contenido_clave", [])
            contexto_old = explanation_data.get("contexto", "")
            insights_old = explanation_data.get("insights", [])
            resumen_old = explanation_data.get("resumen", "")

            # Construimos la explicación didáctica a partir de lo disponible
            explicacion_didactica_new = ""
            if isinstance(resumen_old, str) and resumen_old.strip():
                explicacion_didactica_new = resumen_old.strip()
            else:
                parts = []
                if isinstance(contenido_clave_old, list) and contenido_clave_old:
                    parts.append(" ".join(contenido_clave_old))
                if isinstance(contexto_old, str) and contexto_old.strip():
                    parts.append(contexto_old.strip())
                if isinstance(insights_old, list) and insights_old:
                    parts.append(" ".join(insights_old))
                explicacion_didactica_new = " ".join(p for p in parts if p).strip()

            normalized = {
                "titulo": titulo_old,
                "explicacion_didactica": explicacion_didactica_new or "Explicación generada automáticamente.",
                "puntos_clave": contenido_clave_old if isinstance(contenido_clave_old, list) else [],
                "conexiones": contexto_old if isinstance(contexto_old, str) else "",
                "resumen_corto": resumen_old if isinstance(resumen_old, str) else ""
            }

        return {
            "success": True,
            "slide_number": slide_number,
            "explanation": normalized,
            "raw_response": content
        }

        
    except Exception as e:
        return {
            "success": False,
            "slide_number": slide_number,
            "error": f"Error analyzing slide {slide_number}: {str(e)}"
        }

def generate_word_report(slides: List[bytes], explanations: List[Dict], pdf_name: str) -> bytes:
    """
    Generate a Word document (.docx) report with slides and explanations in continuous format

    Args:
        slides: List of slide image bytes
        explanations: List of explanation dictionaries
        pdf_name: Original PDF name for the report

    Returns:
        Word document bytes
    """
    # Create temporary file for Word document
    tmp_docx_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    tmp_docx_file.close()

    # Keep track of temporary image files to clean up later
    temp_image_files = []

    try:
        # Create Word document
        doc = Document()

        # Set up styles
        title_style = doc.styles.add_style('SlideTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        title_style.font.color.rgb = None  # Default color
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)

        heading_style = doc.styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        heading_style.font.color.rgb = None
        heading_style.paragraph_format.space_after = Pt(6)

        normal_style = doc.styles.add_style('NormalText', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.left_indent = Inches(0.25)
        normal_style.paragraph_format.space_after = Pt(6)

        # Process each slide
        for i, (slide_bytes, explanation) in enumerate(zip(slides, explanations)):
            slide_num = i + 1

            # Slide title
            title_para = doc.add_paragraph(f"Slide {slide_num}", style='SlideTitle')

            # Add slide image
            try:
                # Create temporary image file
                img_tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                img_tmp.write(slide_bytes)
                img_tmp.close()

                temp_image_files.append(img_tmp.name)  # Track for cleanup

                # Add image to document (width: 6 inches, height: auto-maintain aspect ratio)
                doc.add_picture(img_tmp.name, width=Inches(6))

                # Add some space after image
                doc.add_paragraph("")

            except Exception as e:
                error_para = doc.add_paragraph(f"Error loading slide image: {str(e)}", style='NormalText')

            # Add explanation
            if explanation["success"]:
                exp_data = explanation["explanation"]

                title = exp_data.get('titulo') or ""
                exp_did = exp_data.get('explicacion_didactica') or ""
                resumen = exp_data.get('resumen') or ""
                resumen_corto = exp_data.get('resumen_corto') or ""
                puntos = exp_data.get('puntos_clave') or exp_data.get('contenido_clave') or []
                conex = exp_data.get('conexiones') or exp_data.get('contexto') or ""

                explicacion = exp_did if exp_did.strip() else resumen

                if title:
                    doc.add_paragraph("📌 Título", style='SectionHeading')
                    doc.add_paragraph(title, style='NormalText')

                if explicacion:
                    doc.add_paragraph("🧠 Explicación didáctica", style='SectionHeading')
                    doc.add_paragraph(explicacion, style='NormalText')

                if puntos:
                    doc.add_paragraph("🎯 Puntos clave", style='SectionHeading')
                    for item in puntos:
                        doc.add_paragraph(f"• {item}", style='NormalText')

                if conex:
                    doc.add_paragraph("🔗 Conexiones", style='SectionHeading')
                    doc.add_paragraph(conex, style='NormalText')

                # Solo muestra 'Resumen' si es distinto de la explicación
                if resumen and resumen.strip() != explicacion.strip():
                    doc.add_paragraph("📝 Resumen", style='SectionHeading')
                    doc.add_paragraph(resumen, style='NormalText')

                # Solo muestra 'Resumen corto' si es distinto
                if resumen_corto and resumen_corto.strip() not in {resumen.strip(), explicacion.strip()}:
                    doc.add_paragraph("📝 Resumen corto", style='SectionHeading')
                    doc.add_paragraph(resumen_corto, style='NormalText')

            else:
                doc.add_paragraph("❌ Error en el análisis", style='SectionHeading')
                doc.add_paragraph(explanation.get('error', 'Error desconocido'), style='NormalText')

            # Add page break between slides (optional - remove this line for truly continuous)
            # doc.add_page_break()

        # Save the document
        doc.save(tmp_docx_file.name)

        # Read the generated Word document
        with open(tmp_docx_file.name, 'rb') as f:
            docx_bytes = f.read()

        return docx_bytes

    finally:
        # Clean up temporary files
        try:
            os.unlink(tmp_docx_file.name)
        except:
            pass

        for img_file in temp_image_files:
            try:
                os.unlink(img_file)
            except:
                pass

def main():
    """Main Streamlit application"""
    st.set_page_config(
        page_title="PDF Slide Explainer",
        page_icon="📊",
        layout="wide"
    )
    
    st.title("📊 PDF Slide Explainer")
    st.markdown("*Powered by Flora Facturación Infrastructure*")
    
    # Initialize session state
    if 'slides' not in st.session_state:
        st.session_state.slides = None
    if 'explanations' not in st.session_state:
        st.session_state.explanations = None
    if 'edited_explanations' not in st.session_state:
        st.session_state.edited_explanations = None
    if 'uploaded_file_name' not in st.session_state:
        st.session_state.uploaded_file_name = None
    if 'word_report' not in st.session_state:
        st.session_state.word_report = None
    if 'undo_stack' not in st.session_state:
        st.session_state.undo_stack = []
    if 'redo_stack' not in st.session_state:
        st.session_state.redo_stack = []
    if 'current_slide_view' not in st.session_state:
        st.session_state.current_slide_view = None
    if 'selected_language' not in st.session_state:
        st.session_state.selected_language = "Spanish"
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # API Key input
        st.subheader("🔑 OpenAI API Key")
        api_key_input = st.text_input(
            "Enter your OpenAI API Key:",
            type="password",
            help="Your API key will not be stored. Get one at https://platform.openai.com/api-keys"
        )
        
        # Custom prompt
        st.subheader("🎯 Custom Analysis Prompt")
        use_custom_prompt = st.checkbox("Use custom prompt", help="Customize the AI analysis prompt")
        
        custom_prompt = None
        if use_custom_prompt:
            custom_prompt = st.text_area(
                "Custom Prompt:",
                height=200,
                value=PROMPT_TEMPLATE,
                help="Use {slide_number} as placeholder for slide number"
            )
        
        # Language selection
        st.subheader("🌍 Analysis Language")
        languages = [
            "Spanish", "English", "French", "German", "Italian", "Portuguese",
            "Chinese", "Japanese", "Korean", "Arabic", "Russian", "Hindi",
            "Dutch", "Swedish", "Norwegian", "Danish", "Finnish", "Polish",
            "Czech", "Hungarian", "Romanian", "Greek", "Turkish", "Hebrew",
            "Thai", "Vietnamese", "Indonesian", "Malay", "Tagalog", "Swahili"
        ]
        selected_language = st.selectbox(
            "Choose analysis language:",
            options=languages,
            index=0,  # Default to Spanish
            help="Language for AI analysis and explanations"
        )
    
    st.markdown("""
    Upload a PDF presentation and get detailed explanations for each slide using AI vision analysis.
    
    **Features:**
    - 🔍 AI-powered slide analysis using GPT-4 Vision
    - 📋 Structured explanations with key insights
    - 🖼️ Visual element recognition and description
    - 📊 Content extraction and summarization
    - 📄 Professional PDF report generation
    - ✏️ **NEW:** Edit explanations inline
    - 🔍 **NEW:** Enlarge slides for better viewing
    - 🗑️ **NEW:** Delete unwanted slides
    - ↶↷ **NEW:** Undo/Redo functionality (Ctrl+Z / Ctrl+Shift+Z)
    """)
    
    # Initialize OpenAI client
    openai_client = init_openai_client(api_key_input)
    
    if not openai_client:
        if not api_key_input:
            st.warning("⚠️ Please enter your OpenAI API Key in the sidebar to continue.")
        else:
            st.error("❌ Invalid OpenAI API Key. Please check your key.")
        st.stop()
    else:
        st.success("✅ OpenAI client initialized successfully")
    
    # File upload
    st.subheader("📁 Upload PDF Presentation")
    uploaded_file = st.file_uploader(
        "Choose a PDF file",
        type=['pdf'],
        help="Upload a PDF presentation to analyze each slide"
    )
    
    if uploaded_file is not None:
        st.success(f"✅ File uploaded: {uploaded_file.name}")
        
        # Check if this is a new file or same file
        if st.session_state.uploaded_file_name != uploaded_file.name:
            # New file - clear previous results
            st.session_state.slides = None
            st.session_state.explanations = None
            st.session_state.word_report = None
            st.session_state.uploaded_file_name = uploaded_file.name
            
            # Extract slides
            with st.spinner("🔄 Extracting slides from PDF..."):
                st.session_state.slides = extract_slides_from_pdf(uploaded_file)
        
        if not st.session_state.slides:
            st.error("❌ Failed to extract slides from PDF")
            return
        
        st.success(f"✅ Extracted {len(st.session_state.slides)} slides")
        
        # Process slides
        if st.session_state.explanations is None:
            if st.button("🚀 Analyze All Slides", type="primary"):

                # Create progress tracking
                progress_bar = st.progress(0)
                status_text = st.empty()

                explanations = []

                for i, slide_bytes in enumerate(st.session_state.slides):
                    slide_num = i + 1
                    status_text.text(f"Analyzing slide {slide_num} of {len(st.session_state.slides)}...")

                    # Analyze slide
                    explanation = explain_slide(slide_bytes, openai_client, slide_num, custom_prompt, selected_language)
                    explanations.append(explanation)

                    # Update progress
                    progress_bar.progress((i + 1) / len(st.session_state.slides))

                status_text.text("✅ Analysis complete!")

                # Store results in session state
                st.session_state.explanations = explanations
                st.session_state.edited_explanations = [exp.copy() for exp in explanations]  # Initialize edited version
        
        # Display results if available
        if st.session_state.explanations is not None:
            st.subheader("📋 Slide Analysis Results")

            # Undo/Redo buttons
            col1, col2, col3 = st.columns([1, 1, 2])
            with col1:
                if st.button("↶ Undo", disabled=len(st.session_state.undo_stack) == 0):
                    if st.session_state.undo_stack:
                        # Save current state to redo stack
                        st.session_state.redo_stack.append({
                            'slides': st.session_state.slides.copy(),
                            'edited_explanations': [exp.copy() for exp in st.session_state.edited_explanations]
                        })
                        # Restore from undo stack
                        prev_state = st.session_state.undo_stack.pop()
                        st.session_state.slides = prev_state['slides']
                        st.session_state.edited_explanations = prev_state['edited_explanations']
                        st.rerun()

            with col2:
                if st.button("↷ Redo", disabled=len(st.session_state.redo_stack) == 0):
                    if st.session_state.redo_stack:
                        # Save current state to undo stack
                        st.session_state.undo_stack.append({
                            'slides': st.session_state.slides.copy(),
                            'edited_explanations': [exp.copy() for exp in st.session_state.edited_explanations]
                        })
                        # Restore from redo stack
                        next_state = st.session_state.redo_stack.pop()
                        st.session_state.slides = next_state['slides']
                        st.session_state.edited_explanations = next_state['edited_explanations']
                        st.rerun()

            with col3:
                st.caption("💡 Undo/Redo via buttons (keyboard shortcuts not supported in browser)")

            # Use edited explanations if available, otherwise use original
            current_explanations = st.session_state.edited_explanations or st.session_state.explanations

            for i, (slide_bytes, explanation) in enumerate(zip(st.session_state.slides, current_explanations)):
                slide_num = i + 1

                with st.expander(f"📊 Slide {slide_num} Analysis", expanded=True):

                    # Slide controls
                    slide_col1, slide_col2, slide_col3 = st.columns([2, 1, 1])

                    with slide_col1:
                        # Click to enlarge slide
                        if st.button(f"🔍 Enlarge Slide {slide_num}", key=f"enlarge_{i}"):
                            st.session_state.current_slide_view = i
                            st.rerun()

                    with slide_col2:
                        # Edit button
                        if st.button(f"✏️ Edit Text", key=f"edit_{i}"):
                            st.session_state[f"edit_mode_{i}"] = not st.session_state.get(f"edit_mode_{i}", False)
                            st.rerun()

                    with slide_col3:
                        # Delete slide button
                        if st.button(f"🗑️ Delete", key=f"delete_{i}"):
                            # Save current state to undo stack
                            st.session_state.undo_stack.append({
                                'slides': st.session_state.slides.copy(),
                                'edited_explanations': [exp.copy() for exp in current_explanations]
                            })
                            # Remove slide and explanation
                            st.session_state.slides.pop(i)
                            current_explanations.pop(i)
                            st.session_state.edited_explanations = current_explanations
                            # Clear redo stack
                            st.session_state.redo_stack = []
                            st.rerun()

                    # Display slide image (larger, full width)
                    st.subheader(f"🖼️ Slide {slide_num}")
                    st.image(slide_bytes, caption=f"Slide {slide_num}", width='stretch')

                    # Add more space between slide and explanation
                    st.markdown("---")
                    st.markdown("")  # Extra space
                    st.markdown("")  # Extra space

                    # Display explanation below the image
                    st.subheader("🔍 AI Analysis")

                    if explanation["success"]:
                        exp_data = explanation["explanation"]

                        # Check if in edit mode
                        edit_mode = st.session_state.get(f"edit_mode_{i}", False)

                        if edit_mode:
                            # Editable fields
                            st.markdown("### ✏️ Edit Mode")

                            # Title
                            new_title = st.text_input(
                                "📌 Título:",
                                value=exp_data.get('titulo', ''),
                                key=f"title_{i}"
                            )

                            # Explicacion didactica
                            new_explicacion_didactica = st.text_area(
                                "🧠 Explicación didáctica:",
                                value=exp_data.get('explicacion_didactica', ''),
                                key=f"explicacion_{i}",
                                height=100
                            )

                            # Puntos clave
                            st.markdown("**🎯 Puntos clave:**")
                            new_puntos_clave = []
                            if exp_data.get('puntos_clave'):
                                for j, item in enumerate(exp_data['puntos_clave']):
                                    new_item = st.text_input(
                                        f"Item {j+1}:",
                                        value=item,
                                        key=f"punto_{i}_{j}"
                                    )
                                    new_puntos_clave.append(new_item)

                            # Conexiones
                            new_conexiones = st.text_area(
                                "🔗 Conexiones:",
                                value=exp_data.get('conexiones', ''),
                                key=f"conexiones_{i}",
                                height=80
                            )

                            # Resumen corto
                            new_resumen_corto = st.text_area(
                                "📝 Resumen corto:",
                                value=exp_data.get('resumen_corto', ''),
                                key=f"resumen_corto_{i}",
                                height=60
                            )

                            # Save button
                            if st.button("💾 Save Changes", key=f"save_{i}"):
                                # Save current state to undo stack
                                st.session_state.undo_stack.append({
                                    'slides': st.session_state.slides.copy(),
                                    'edited_explanations': [exp.copy() for exp in current_explanations]
                                })

                                # Update the explanation with new schema
                                updated_exp = explanation.copy()
                                updated_exp["explanation"] = {
                                    'titulo': new_title,
                                    'explicacion_didactica': new_explicacion_didactica,
                                    'puntos_clave': new_puntos_clave,
                                    'conexiones': new_conexiones,
                                    'resumen_corto': new_resumen_corto
                                }

                                # Update edited explanations
                                if st.session_state.edited_explanations is None:
                                    st.session_state.edited_explanations = [exp.copy() for exp in st.session_state.explanations]
                                st.session_state.edited_explanations[i] = updated_exp

                                # Clear redo stack
                                st.session_state.redo_stack = []

                                # Exit edit mode
                                st.session_state[f"edit_mode_{i}"] = False
                                st.success("✅ Changes saved!")
                                st.rerun()

                        else:
                            # Display mode (nuevo esquema con fallback)
                            st.markdown(f"**📌 Título:** {exp_data.get('titulo', 'N/A')}")

                            explicacion_ui = exp_data.get('explicacion_didactica') or exp_data.get('resumen') or 'N/A'
                            st.markdown("**🧠 Explicación didáctica:**")
                            st.write(explicacion_ui)

                            puntos_ui = exp_data.get('puntos_clave') or exp_data.get('contenido_clave') or []
                            if puntos_ui:
                                st.markdown("**🎯 Puntos clave:**")
                                for item in puntos_ui:
                                    st.markdown(f"- {item}")

                            conex_ui = exp_data.get('conexiones') or exp_data.get('contexto') or ''
                            if conex_ui:
                                st.markdown("**🔗 Conexiones:**")
                                st.write(conex_ui)

                            resumen_corto_ui = exp_data.get('resumen_corto') or exp_data.get('resumen') or ''
                            if resumen_corto_ui:
                                st.markdown("**📝 Resumen corto:**")
                                st.write(resumen_corto_ui)

                            insights_ui = exp_data.get('insights') or []


                    else:
                        st.error(f"❌ {explanation.get('error', 'Unknown error')}")

            # Modal for enlarged slide view
            if st.session_state.current_slide_view is not None:
                i = st.session_state.current_slide_view
                slide_bytes = st.session_state.slides[i]
                slide_num = i + 1

                with st.container():
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.image(slide_bytes, caption=f"Slide {slide_num} (Enlarged)", use_container_width=True)
                    with col2:
                        if st.button("❌ Close", key="close_enlarged"):
                            st.session_state.current_slide_view = None
                            st.rerun()
            
            # Export options
            st.subheader("📤 Export Results")
            
            col1, col2 = st.columns(2)
            
            with col1:
                # JSON Export
                export_data = {
                    "pdf_name": st.session_state.uploaded_file_name,
                    "total_slides": len(st.session_state.slides),
                    "analysis_timestamp": str(datetime.now()),
                    "explanations": st.session_state.explanations
                }
                
                export_json = json.dumps(export_data, indent=2, ensure_ascii=False)
                
                st.download_button(
                    label="📥 Download Analysis as JSON",
                    data=export_json,
                    file_name=f"{st.session_state.uploaded_file_name.replace('.pdf', '')}_analysis.json",
                    mime="application/json",
                    key="json_download"
                )
            
            with col2:
                # Word Document Export
                if st.button("📄 Generate Word Report", key="generate_word"):
                    with st.spinner("🔄 Generating Word document... This may take a moment depending on the number of slides."):
                        try:
                            # Use edited explanations if available, otherwise use original
                            explanations_to_use = st.session_state.edited_explanations or st.session_state.explanations
                            st.session_state.word_report = generate_word_report(
                                st.session_state.slides,
                                explanations_to_use,
                                st.session_state.uploaded_file_name
                            )
                            st.success("✅ Word document generated successfully!")

                        except Exception as e:
                            st.error(f"❌ Error generating Word document: {str(e)}")

                # Show Word download button if report is generated
                if st.session_state.word_report is not None:
                    st.download_button(
                        label="📥 Download Word Report",
                        data=st.session_state.word_report,
                        file_name=f"{st.session_state.uploaded_file_name.replace('.pdf', '')}_analysis_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="word_download"
                    )

if __name__ == "__main__":
    main()